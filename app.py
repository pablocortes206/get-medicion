from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime, timedelta
from typing import Optional, Dict, List, Tuple
import io
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

import pandas as pd
import subprocess
import tempfile
import os
import streamlit as st
from supabase import create_client, Client

# =========================================================
# CONFIG
# =========================================================
st.set_page_config(page_title="GET Wear Monitor", layout="wide")

APP_VERSION = "9.5"

def enforce_version():
    qp = st.query_params.to_dict()
    version_url = qp.get("v", "")
    version_session = st.session_state.get("app_version", "")

    if version_url != APP_VERSION:
        st.query_params["v"] = APP_VERSION
        st.rerun()

    if version_session != APP_VERSION:
        st.session_state["app_version"] = APP_VERSION
        st.cache_data.clear()

enforce_version()

# ── PALETA TECK ────────────────────────────────────────────
TECK_GREEN   = "#007A3D"
TECK_GREEN_2 = "#00A04A"
TECK_DARK    = "#0B0F14"

# ── PALETA TECH LEGIBLE (fondo gris claro) ─────────────────
# Fondos de tarjeta por estado
BG_OK     = "#F4FAF0"   # verde muy claro
BG_MEDIO  = "#FFFBF2"   # ámbar muy claro
BG_ALTO   = "#FFF3EC"   # naranja muy claro
BG_CRIT   = "#FFF0F0"   # rojo muy claro
BG_NODAT  = "#F5F5F5"   # gris neutro

# Bordes por estado
BORDER_OK    = "#C0DD97"
BORDER_MEDIO = "#FAC775"
BORDER_ALTO  = "#F0997B"
BORDER_CRIT  = "#F09595"

# Texto por estado (oscuro del mismo tono)
TEXT_OK    = "#27500A"
TEXT_MEDIO = "#633806"
TEXT_ALTO  = "#712B13"
TEXT_CRIT  = "#791F1F"

# Badge % desgaste
PCT_LO_BG  = "#EAF3DE"; PCT_LO_FG  = "#27500A"   # 0-44%
PCT_MD_BG  = "#FAEEDA"; PCT_MD_FG  = "#633806"   # 45-74%
PCT_HI_BG  = "#FCEBEB"; PCT_HI_FG  = "#791F1F"   # 75-89%
PCT_CR_BG  = "#A32D2D"; PCT_CR_FG  = "#FFFFFF"   # 90%+

# Badge días sin medir
DIAS_OK_BG   = "#EAF3DE"; DIAS_OK_FG   = "#27500A"
DIAS_MD_BG   = "#FAEEDA"; DIAS_MD_FG   = "#633806"
DIAS_CR_BG   = "#A32D2D"; DIAS_CR_FG   = "#FFFFFF"

EQUIPOS = [
    "101","102","103","104","105","106","108",
    "201","202","203","204","205",
    "301","302","303",
]
EQUIPOS_MOTONIVELADORA = {"301","302","303"}

CODIGO_A_ID = {
    "0174-DZ-101":"101","0174-DZ-102":"102","0174-DZ-103":"103",
    "0174-DZ-104":"104","0174-DZ-105":"105","0174-DZ-106":"106",
    "0174-DZ-108":"108",
    "0174-WD-201":"201","0174-WD-202":"202","0174-WD-203":"203",
    "0174-WD-204":"204","0174-WD-205":"205",
    "0174-GR-301":"301","0174-GR-302":"302","0174-GR-303":"303",
}

REGLAS: Dict[str, dict] = {
    "MOTONIVELADORA": {
        "puntos": [(122,100),(145,91),(167,78),(190,65),(212,52),(235,39),(257,26),(280,13),(302,0)],
        "umbrales": [("CRÍTICO",90,"Cambiar cuchilla."),("MEDIO",65,"Monitorear."),("OK",0,"Operación normal.")],
        "mm_nuevo": 302.0, "mm_critico": 145.0,
    },
    "DOZER_854_D10_D11": {
        "puntos": [(75,100),(82,95),(83,90),(100,75),(140,45),(170,0)],
        "umbrales": [("CRÍTICO",95,"Detención inmediata."),("ALTO",75,"Programar cambio."),("MEDIO",45,"Monitorear."),("OK",0,"Operación normal.")],
        "mm_nuevo": 170.0, "mm_critico": 82.0,
    },
}

COLOR_ESTADO = {"OK":"🟢","MEDIO":"🟡","ALTO":"🟠","CRÍTICO":"🔴"}

# Mapas de color para tarjetas (bg fila, border, text)
_ESTADO_CARD = {
    "OK":     (BG_OK,    BORDER_OK,    TEXT_OK),
    "MEDIO":  (BG_MEDIO, BORDER_MEDIO, TEXT_MEDIO),
    "ALTO":   (BG_ALTO,  BORDER_ALTO,  TEXT_ALTO),
    "CRÍTICO":(BG_CRIT,  BORDER_CRIT,  TEXT_CRIT),
}

# Mantener BG_ESTADO para compatibilidad con código que lo usa
BG_ESTADO = {
    "OK":     BG_OK,
    "MEDIO":  BG_MEDIO,
    "ALTO":   BG_ALTO,
    "CRÍTICO":BG_CRIT,
}

HORO_TOLERANCIA_PCT = 5.0


# =========================================================
# HELPERS DE COLOR (TECH)
# =========================================================

def badge_pct_html(pct_val) -> str:
    """Devuelve HTML de badge de porcentaje con paleta tech legible."""
    if pct_val is None or not str(pct_val).replace('.','').replace('-','').isdigit():
        try:
            p = float(pct_val)
        except:
            return "—"
    else:
        p = float(pct_val)

    if p >= 90:
        bg, fg = PCT_CR_BG, PCT_CR_FG
    elif p >= 75:
        bg, fg = PCT_HI_BG, PCT_HI_FG
    elif p >= 45:
        bg, fg = PCT_MD_BG, PCT_MD_FG
    else:
        bg, fg = PCT_LO_BG, PCT_LO_FG
    return f'<span style="background:{bg};color:{fg};padding:2px 10px;border-radius:6px;font-weight:bold;">{p:.1f}%</span>'


def badge_dias_html(dias: int) -> str:
    """Devuelve HTML de badge de días sin medir con paleta tech legible."""
    if dias <= 10:
        bg, fg = DIAS_OK_BG, DIAS_OK_FG
    elif dias <= 14:
        bg, fg = DIAS_MD_BG, DIAS_MD_FG
    else:
        bg, fg = DIAS_CR_BG, DIAS_CR_FG
    return f'<span style="background:{bg};color:{fg};padding:2px 10px;border-radius:6px;font-weight:bold;">{dias} días sin medir</span>'


def card_style(estado: str) -> str:
    """Devuelve el style inline de una tarjeta equipo."""
    bg, border, _ = _ESTADO_CARD.get(estado, (BG_NODAT, "#D3D1C7", "#444441"))
    return (
        f"background:{bg};"
        f"border:1px solid {border};"
        f"border-radius:10px;"
        f"padding:10px 16px;"
        f"margin-bottom:6px;"
    )


def dot_estado(estado: str) -> str:
    """Dot de color sólido según estado."""
    colors = {"OK":"#3B6D11","MEDIO":"#BA7517","ALTO":"#854F0B","CRÍTICO":"#A32D2D"}
    c = colors.get(estado, "#888780")
    return f'<span style="display:inline-block;width:9px;height:9px;border-radius:50%;background:{c};margin-right:6px;vertical-align:middle;"></span>'


# =========================================================
# ESTILO GLOBAL
# =========================================================
def inject_style():
    st.markdown(f"""
    <style>
    /* Fondo principal: gris claro tech */
    .stApp {{
        background: #F1EFE8;
    }}
    .block-container {{
        padding-top:1rem !important;
        padding-bottom:2rem !important;
    }}
    /* Header con acento verde Teck sobre gris claro */
    .teck-header {{
        display:flex;
        align-items:center;
        justify-content:space-between;
        gap:1rem;
        padding:16px 20px;
        border-radius:16px;
        background:#FFFFFF;
        border:1px solid {BORDER_OK};
        border-left:5px solid {TECK_GREEN};
        box-shadow:0 2px 8px rgba(0,0,0,0.08);
        margin-bottom:18px;
    }}
    .teck-header p {{
        color:#2C2C2A !important;
    }}
    .teck-badge {{
        padding:8px 14px;
        border-radius:999px;
        font-size:.85rem;
        font-weight:800;
        color:white;
        background:{TECK_GREEN};
        border:1px solid {TECK_GREEN_2};
        white-space:nowrap;
    }}
    /* Botones */
    div.stButton>button {{
        border-radius:10px !important;
        font-weight:700 !important;
        color:#2C2C2A !important;
        background:#FFFFFF !important;
        border:1px solid #D3D1C7 !important;
    }}
    div.stButton>button[kind="primary"] {{
        background:{TECK_GREEN} !important;
        color:white !important;
        border-color:{TECK_GREEN_2} !important;
    }}
    div.stButton>button[kind="primary"]:hover {{
        background:{TECK_GREEN_2} !important;
    }}
    /* Tarjetas equipo */
    .equipo-card {{
        padding:10px 16px;
        border-radius:10px;
        margin-bottom:6px;
    }}
    /* Texto general más oscuro para legibilidad sobre gris claro */
    .stMarkdown, .stText, label, .stSelectbox label, .stNumberInput label {{
        color:#2C2C2A !important;
    }}
    /* Métricas */
    [data-testid="stMetric"] {{
        background:#FFFFFF;
        border:1px solid #D3D1C7;
        border-radius:10px;
        padding:10px 14px;
    }}
    /* Sidebar */
    [data-testid="stSidebar"] {{
        background:#FFFFFF;
        border-right:1px solid #D3D1C7;
    }}
    /* Tabs */
    .stTabs [data-baseweb="tab-list"] {{
        background:#FFFFFF;
        border-radius:10px;
        border:1px solid #D3D1C7;
        padding:4px;
    }}
    .stTabs [data-baseweb="tab"] {{
        color:#5F5E5A !important;
        font-weight:600;
    }}
    .stTabs [aria-selected="true"] {{
        background:{TECK_GREEN} !important;
        color:white !important;
        border-radius:8px;
    }}
    /* Dataframes */
    [data-testid="stDataFrame"] {{
        border:1px solid #D3D1C7;
        border-radius:10px;
    }}
    /* Divider */
    hr {{
        border-color:#D3D1C7 !important;
    }}
    /* Info / Warning / Error boxes */
    .stAlert {{
        border-radius:10px !important;
    }}
    </style>
    """, unsafe_allow_html=True)


def render_header():
    st.markdown("""
    <div class="teck-header">
      <div>
        <p style="font-size:42px;font-weight:900;margin:0;line-height:1.05;color:#0B1E0F !important;">GET Wear Monitor</p>
        <p style="font-size:18px;margin:6px 0 0 0;color:#3B6D11 !important;">Sistema de monitoreo y proyección de desgaste de cuchillas</p>
        <p style="font-size:14px;margin:8px 0 0 0;color:#5F5E5A !important;"><b>Creado por:</b> Pablo Cortés Ramos · Ingeniero de Mantenimiento / Confiabilidad &nbsp;|&nbsp; <b style="color:#007A3D;">v9.5 — 04/05/2026</b></p>
      </div>
      <div class="teck-badge">Teck QB2 · GET Wear Monitor</div>
    </div>
    """, unsafe_allow_html=True)


# =========================================================
# INICIALIZACIÓN
# =========================================================
inject_style()
render_header()


# =========================================================
# SUPABASE
# =========================================================
@st.cache_resource
def get_supabase() -> Client:
    return create_client(st.secrets["SUPABASE_URL"].strip(), st.secrets["SUPABASE_KEY"].strip())

def sb() -> Client:
    return get_supabase()


# =========================================================
# LÓGICA
# =========================================================
def regla_por_equipo(eq: str) -> str:
    return "MOTONIVELADORA" if eq in EQUIPOS_MOTONIVELADORA else "DOZER_854_D10_D11"

def rango_regla(regla: str) -> tuple[float,float]:
    xs = [p[0] for p in REGLAS[regla]["puntos"]]
    return min(xs), max(xs)

def interpolar_pct(mm: float, puntos) -> float:
    puntos = sorted(puntos, key=lambda x: x[0])
    if mm <= puntos[0][0]: return float(puntos[0][1])
    if mm >= puntos[-1][0]: return float(puntos[-1][1])
    for (x1,y1),(x2,y2) in zip(puntos[:-1],puntos[1:]):
        if x1 <= mm <= x2:
            return float(y1 + (mm-x1)/(x2-x1)*(y2-y1))
    return float(puntos[-1][1])

def clasificar(pct, umbrales):
    for estado,limite,accion in umbrales:
        if pct >= limite: return estado, accion
    return "OK","Operación normal."

def calcular_tasa_meds(meds: list[dict]) -> Optional[float]:
    if len(meds) < 2: return None
    tasas = []
    for i in range(len(meds)-1):
        dh  = float(meds[i]["horometro"]) - float(meds[i+1]["horometro"])
        dmm = float(meds[i+1]["mm_usada"]) - float(meds[i]["mm_usada"])
        if dh > 0 and dmm > 0: tasas.append(dmm/dh)
    return round(sum(tasas)/len(tasas),4) if tasas else None

def proyectar(mm_usada: float, tasa: Optional[float], mm_critico: float):
    if not tasa or tasa <= 0: return None, None
    r = mm_usada - mm_critico
    if r <= 0: return 0.0, 0.0
    h = r / tasa
    return round(h,1), round(h/24,1)


# =========================================================
# LECTURA EXCEL HORÓMETROS
# =========================================================
def leer_excel_horometros(archivo) -> pd.DataFrame:
    from openpyxl import load_workbook
    wb = load_workbook(archivo, read_only=True, data_only=True)
    ws = wb["HOROMETRO"]

    registros = []
    for row in ws.iter_rows(min_row=29, max_row=43, min_col=2, max_col=8, values_only=True):
        codigo = row[0]
        if not codigo or not isinstance(codigo, str): continue
        codigo = codigo.strip()
        if codigo not in CODIGO_A_ID: continue

        equipo_id = CODIGO_A_ID[codigo]
        fecha_act = row[2]
        if hasattr(fecha_act, 'date'):
            fecha_act = fecha_act.date()
        elif isinstance(fecha_act, str):
            try: fecha_act = datetime.strptime(fecha_act, "%Y-%m-%d").date()
            except: fecha_act = date.today()
        else:
            fecha_act = date.today()

        horometro   = float(row[3]) if row[3] is not None else None
        prom_7d     = float(row[4]) if row[4] is not None else None
        prom_30d    = float(row[5]) if row[5] is not None else None
        prom_hist   = float(row[6]) if row[6] is not None else None

        if horometro is None: continue

        def corregir_promedio(val, hist_ref):
            if val is None: return None
            if val < 0 or val > 24:
                return hist_ref if (hist_ref and 0 < hist_ref <= 24) else None
            return val

        prom_7d   = corregir_promedio(prom_7d,   prom_hist)
        prom_30d  = corregir_promedio(prom_30d,  prom_hist)
        prom_hist = corregir_promedio(prom_hist, None)

        registros.append({
            "codigo_excel": codigo,
            "equipo": equipo_id,
            "fecha": str(fecha_act),
            "horometro_actual": horometro,
            "promedio_7d": prom_7d,
            "promedio_30d": prom_30d,
            "promedio_historico": prom_hist,
            "tiene_error": (row[4] is not None and float(row[4]) < 0) or (row[5] is not None and float(row[5]) < 0),
        })

    wb.close()
    return pd.DataFrame(registros)


def guardar_horometros(df: pd.DataFrame) -> int:
    if df.empty: return 0
    count = 0
    for _, row in df.iterrows():
        payload = {
            "fecha": row["fecha"],
            "equipo": row["equipo"],
            "horometro_actual": row["horometro_actual"],
            "promedio_7d": row["promedio_7d"],
            "promedio_30d": row["promedio_30d"],
            "promedio_historico": row["promedio_historico"],
            "creado_en": datetime.utcnow().isoformat(),
        }
        sb().table("horometros").insert(payload).execute()
        count += 1
    cargar_horometros_db.clear()
    return count


@st.cache_data(ttl=300)
def cargar_horometros_db() -> pd.DataFrame:
    try:
        resp = sb().table("horometros").select("*").order("fecha", desc=True).limit(500).execute()
        df = pd.DataFrame(resp.data or [])
        if df.empty: return df
        df["fecha_dt"] = pd.to_datetime(df["fecha"], errors="coerce")
        return df.sort_values("fecha_dt", ascending=False).drop_duplicates("equipo", keep="first")
    except Exception:
        return pd.DataFrame()


def proyectar_con_horas_dia(mm_usada, h_dia, mm_critico, horometro_actual, equipo):
    try:
        resp = sb().table("mediciones").select("horometro,mm_usada").eq("equipo", equipo).eq("es_cambio", False).order("fecha", desc=True).limit(10).execute()
        meds_all = resp.data or []
    except:
        meds_all = []
    tasa = calcular_tasa_meds(meds_all)
    if tasa is None or tasa <= 0:
        regla = regla_por_equipo(equipo)
        tasa = 0.013 if regla == "DOZER_854_D10_D11" else 0.028
    restante_mm = mm_usada - mm_critico
    if restante_mm <= 0: return 0.0, 0.0
    horas = restante_mm / tasa
    dias  = horas / h_dia
    return round(horas, 1), round(dias, 1)


def evaluar(eq, horometro, mm_izq, mm_der) -> dict:
    mm_usada = min(mm_izq, mm_der)
    regla = regla_por_equipo(eq)
    cfg = REGLAS[regla]
    pct = round(interpolar_pct(mm_usada, cfg["puntos"]), 1)
    estado, accion = clasificar(pct, cfg["umbrales"])
    hist = ultimas_meds_equipo(eq, 5)
    meds = [{"horometro": horometro, "mm_usada": mm_usada}] + hist
    tasa = calcular_tasa_meds(meds)
    if tasa is None:
        df_horo = cargar_horometros_db()
        if not df_horo.empty:
            row_h = df_horo[df_horo["equipo"] == eq]
            if not row_h.empty:
                prom = row_h.iloc[0].get("promedio_30d") or row_h.iloc[0].get("promedio_historico")
                if prom and prom > 0:
                    h_c, d_c = proyectar_con_horas_dia(mm_usada, float(prom), cfg["mm_critico"], horometro, eq)
                    return dict(mm_usada=mm_usada, pct=pct, estado=estado, accion=accion,
                                tasa=None, h_critico=h_c, d_critico=d_c, regla=regla,
                                fuente_tasa="Excel horómetros")
    h_c, d_c = proyectar(mm_usada, tasa, cfg["mm_critico"])
    return dict(mm_usada=mm_usada, pct=pct, estado=estado, accion=accion,
                tasa=tasa, h_critico=h_c, d_critico=d_c, regla=regla,
                fuente_tasa="Mediciones reales" if tasa else "Sin datos")


# =========================================================
# DB
# =========================================================
@st.cache_data(ttl=60)
def cargar_historial(limit: int = 500) -> pd.DataFrame:
    try:
        resp = sb().table("mediciones").select("*").order("fecha",desc=True).limit(limit).execute()
        return pd.DataFrame(resp.data or [])
    except Exception:
        return pd.DataFrame()

@st.cache_data(ttl=60)
def cargar_cambios(limit: int = 200) -> pd.DataFrame:
    try:
        resp = sb().table("cambios_cuchilla").select("*").order("fecha",desc=True).limit(limit).execute()
        return pd.DataFrame(resp.data or [])
    except Exception:
        return pd.DataFrame()

def ultimo_cambio_equipo(eq: str) -> Optional[dict]:
    try:
        resp = sb().table("cambios_cuchilla").select("fecha,horometro").eq("equipo",eq).order("fecha",desc=True).limit(1).execute()
        data = resp.data or []
        return data[0] if data else None
    except Exception:
        return None

def ultimas_meds_equipo(eq: str, n: int = 5) -> list[dict]:
    try:
        uc = ultimo_cambio_equipo(eq)
        q = sb().table("mediciones").select("fecha,horometro,mm_usada").eq("equipo",eq).eq("es_cambio",False).order("fecha",desc=True).limit(n)
        if uc: q = q.gte("fecha", uc["fecha"])
        return q.execute().data or []
    except Exception:
        return []

def guardar_medicion(fecha, eq, horometro, mm_izq, mm_der, usuario, r, sospechoso=False, comentario_sospecha=""):
    sb().table("mediciones").insert({
        "fecha": str(fecha), "equipo": eq,
        "horometro": float(horometro), "mm_izq": float(mm_izq), "mm_der": float(mm_der),
        "mm_usada": float(r["mm_usada"]), "condicion_pct": float(r["pct"]),
        "estado": r["estado"], "accion": r["accion"],
        "tasa_mm_h": float(r["tasa"]) if r.get("tasa") else None,
        "horas_a_critico": float(r["h_critico"]) if r.get("h_critico") is not None else None,
        "dias_a_critico": float(r["d_critico"]) if r.get("d_critico") is not None else None,
        "usuario": usuario.strip(), "componente": "Cuchilla", "es_cambio": False,
        "sospechoso": sospechoso,
        "comentario_sospecha": comentario_sospecha if sospechoso else "",
        "creado_en": datetime.utcnow().isoformat(),
    }).execute()
    cargar_historial.clear()

def marcar_sospechosa(id_med: int, sospechoso: bool, comentario: str = ""):
    sb().table("mediciones").update({
        "sospechoso": sospechoso,
        "comentario_sospecha": comentario if sospechoso else "",
    }).eq("id", int(id_med)).execute()
    cargar_historial.clear()

def eliminar_medicion(id_med: int):
    sb().table("mediciones").delete().eq("id", int(id_med)).execute()
    cargar_historial.clear()

UMBRAL_SOSPECHA_MM = 10.0

def verificar_sospecha(equipo: str, mm_nueva: float) -> dict:
    try:
        uc = ultimo_cambio_equipo(equipo)
        q = sb().table("mediciones").select("fecha,mm_usada,horometro").eq("equipo", equipo).eq("es_cambio", False).order("fecha", desc=True).limit(1)
        if uc:
            q = q.gte("fecha", uc["fecha"])
        data = q.execute().data or []
        if not data:
            return {"sospechoso": False}
        mm_anterior = float(data[0]["mm_usada"])
        diferencia = mm_nueva - mm_anterior
        if diferencia >= UMBRAL_SOSPECHA_MM:
            return {
                "sospechoso": True,
                "mm_anterior": mm_anterior,
                "mm_nueva": mm_nueva,
                "diferencia": round(diferencia, 1),
                "comentario_auto": f"Subió {diferencia:.1f} mm ({mm_anterior:.1f} → {mm_nueva:.1f} mm) sin cambio registrado.",
            }
        return {"sospechoso": False}
    except Exception:
        return {"sospechoso": False}

def actualizar_medicion(id_med: int, horometro: float, mm_izq: float, mm_der: float, usuario_edit: str, eq: str):
    r = evaluar(eq, horometro, mm_izq, mm_der)
    sb().table("mediciones").update({
        "horometro": float(horometro),
        "mm_izq": float(mm_izq),
        "mm_der": float(mm_der),
        "mm_usada": float(r["mm_usada"]),
        "condicion_pct": float(r["pct"]),
        "estado": r["estado"],
        "accion": r["accion"],
        "tasa_mm_h": float(r["tasa"]) if r.get("tasa") else None,
        "horas_a_critico": float(r["h_critico"]) if r.get("h_critico") is not None else None,
        "dias_a_critico": float(r["d_critico"]) if r.get("d_critico") is not None else None,
        "usuario": usuario_edit.strip(),
    }).eq("id", int(id_med)).execute()
    cargar_historial.clear()

def guardar_cambio(fecha, eq, horometro, mm_izq_f, mm_der_f, fue_virada, motivo, obs, tec1, tec2, usuario):
    sb().table("cambios_cuchilla").insert({
        "fecha": str(fecha), "equipo": eq, "horometro": float(horometro),
        "mm_izq_final": float(mm_izq_f), "mm_der_final": float(mm_der_f),
        "fue_virada": fue_virada, "motivo": motivo,
        "observaciones": obs.strip() if obs else None,
        "tecnico_1": tec1.strip(), "tecnico_2": tec2.strip() if tec2 else None,
        "usuario": usuario.strip(), "creado_en": datetime.utcnow().isoformat(),
    }).execute()
    regla = regla_por_equipo(eq)
    mm_nuevo = REGLAS[regla]["mm_nuevo"]
    sb().table("mediciones").insert({
        "fecha": str(fecha), "equipo": eq, "horometro": float(horometro),
        "mm_izq": mm_nuevo, "mm_der": mm_nuevo, "mm_usada": mm_nuevo,
        "condicion_pct": 0.0, "estado": "OK", "accion": "GET nuevo instalado.",
        "usuario": tec1.strip() if tec1 else usuario.strip(),
        "componente": "Cuchilla NUEVA", "es_cambio": True,
        "creado_en": datetime.utcnow().isoformat(),
    }).execute()
    cargar_historial.clear()
    cargar_cambios.clear()

def ultimos_por_equipo(df: pd.DataFrame) -> pd.DataFrame:
    if df.empty: return df
    df2 = df.copy()
    df2["fecha_dt"] = pd.to_datetime(df2["fecha"], errors="coerce")
    return df2.sort_values("fecha_dt", ascending=False).drop_duplicates("equipo", keep="first")

def excel_bytes(df: pd.DataFrame) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        df.to_excel(w, index=False, sheet_name="Datos")
    return buf.getvalue()

def dias_sin_medicion(df: pd.DataFrame) -> pd.DataFrame:
    hoy = date.today()
    rows = []
    for eq in EQUIPOS:
        df_eq = df[df["equipo"] == eq] if not df.empty else pd.DataFrame()
        if df_eq.empty:
            rows.append({"equipo": eq, "ultima_medicion": "—", "dias_sin_medir": None, "estado": "SIN DATOS"})
        else:
            ultima = pd.to_datetime(df_eq["fecha"]).max().date()
            dias = (hoy - ultima).days
            estado_last = df_eq.sort_values("fecha", ascending=False).iloc[0].get("estado","?")
            rows.append({"equipo": eq, "ultima_medicion": str(ultima), "dias_sin_medir": dias, "estado": estado_last})
    return pd.DataFrame(rows)

def color_dias(dias) -> str:
    """Color de texto para días sin medir (compatible con código legacy)."""
    if dias is None: return "#888780"
    if dias <= 10:   return "#3B6D11"
    if dias <= 14:   return "#BA7517"
    return "#A32D2D"


# =========================================================
# VALIDACIÓN HORÓMETRO
# =========================================================
def validar_horometro(equipo: str, horometro_ingresado: float) -> dict:
    df_horo = cargar_horometros_db()
    if df_horo.empty: return {"ok": True}
    row = df_horo[df_horo["equipo"] == equipo]
    if row.empty: return {"ok": True}
    h_excel = float(row.iloc[0]["horometro_actual"])
    fecha_excel = row.iloc[0]["fecha"]
    if h_excel <= 0: return {"ok": True}
    diferencia_pct = abs(horometro_ingresado - h_excel) / h_excel * 100
    if diferencia_pct > HORO_TOLERANCIA_PCT:
        return {
            "ok": False,
            "h_excel": h_excel,
            "fecha_excel": fecha_excel,
            "diferencia": abs(horometro_ingresado - h_excel),
            "diferencia_pct": round(diferencia_pct, 1),
        }
    return {"ok": True}


# =========================================================
# REPORTE CORREO
# =========================================================
def color_estado_html(estado: str) -> str:
    m = {"OK": PCT_LO_BG, "MEDIO": PCT_MD_BG, "ALTO":"#FAEEDA","CRÍTICO": PCT_CR_BG,"SIN DATOS":"#F1EFE8"}
    return m.get(estado, "#F1EFE8")

def color_estado_text_html(estado: str) -> str:
    m = {"OK": TEXT_OK, "MEDIO": TEXT_MEDIO, "ALTO": TEXT_ALTO, "CRÍTICO":"#FFFFFF","SIN DATOS":"#888780"}
    return m.get(estado, "#444441")

def texto_estado_html(estado: str) -> str:
    m = {"OK":"✅ OK","MEDIO":"🟡 Monitoreo","ALTO":"🟠 Programar cambio","CRÍTICO":"🔴 CRÍTICO","SIN DATOS":"⚫ Sin datos"}
    return m.get(estado, estado)

def color_dias_html(dias) -> str:
    if dias is None: return "#F1EFE8"
    if dias <= 10:   return PCT_LO_BG
    if dias <= 14:   return PCT_MD_BG
    return PCT_CR_BG

def color_dias_text_html(dias) -> str:
    if dias is None: return "#888780"
    if dias <= 10:   return TEXT_OK
    if dias <= 14:   return TEXT_MEDIO
    return "#FFFFFF"

def generar_html_reporte(df_estados: pd.DataFrame, df_sin_medir: pd.DataFrame) -> str:
    filas = ""
    if not df_estados.empty:
        for _, r in df_estados.iterrows():
            est = str(r.get("estado",""))
            bg_est  = color_estado_html(est)
            fg_est  = color_estado_text_html(est)
            txt_est = texto_estado_html(est)
            mm  = f"{r['mm_usada']:.1f}" if pd.notna(r.get("mm_usada")) else "—"
            pct = f"{r['condicion_pct']:.1f}%" if pd.notna(r.get("condicion_pct")) else "—"
            dias_c_val = r.get("dias_a_critico")
            if pd.notna(dias_c_val):
                dias_c_num = float(dias_c_val)
                bg_dc = color_dias_html(dias_c_num)
                fg_dc = color_dias_text_html(dias_c_num)
                dias_c = f'<span style="background:{bg_dc};color:{fg_dc};padding:2px 8px;border-radius:4px;">{dias_c_num:.0f} días</span>'
            else:
                dias_c = "—"
            filas += f"""<tr>
              <td style="padding:8px;font-weight:bold;color:#2C2C2A;">{r['equipo']}</td>
              <td style="padding:8px;color:#444441;">{r.get('fecha','—')}</td>
              <td style="padding:8px;color:#444441;">{mm} mm</td>
              <td style="padding:8px;color:#444441;">{pct}</td>
              <td style="padding:8px;background:{bg_est};color:{fg_est};border-radius:6px;text-align:center;font-weight:bold;">{txt_est}</td>
              <td style="padding:8px;">{dias_c}</td>
            </tr>"""

    sin_medir = df_sin_medir[df_sin_medir["dias_sin_medir"].notna() & (df_sin_medir["dias_sin_medir"] > 7)] if not df_sin_medir.empty else pd.DataFrame()
    filas_sm = ""
    for _, r in sin_medir.iterrows():
        d = int(r['dias_sin_medir'])
        bg = color_dias_html(d)
        fg = color_dias_text_html(d)
        filas_sm += f"<tr><td style='padding:8px;font-weight:bold;color:#2C2C2A;'>{r['equipo']}</td><td style='padding:8px;color:#444441;'>{r.get('ultima_medicion','—')}</td><td style='padding:8px;'><span style='background:{bg};color:{fg};padding:2px 8px;border-radius:4px;font-weight:bold;'>{d} días</span></td></tr>"

    tabla_sm = f"""<h2 style="color:#A32D2D;margin-top:30px;">⚠️ Sin medir hace más de 7 días</h2>
    <table style="width:100%;border-collapse:collapse;background:#FFFFFF;border-radius:8px;border:1px solid #D3D1C7;">
    <thead><tr style="background:{TECK_GREEN};color:white;"><th style="padding:10px;text-align:left;">Equipo</th>
    <th style="padding:10px;text-align:left;">Última medición</th><th style="padding:10px;text-align:left;">Días sin medir</th>
    </tr></thead><tbody>{filas_sm}</tbody></table>""" if filas_sm else "<p style='color:#27500A;font-weight:bold;'>✅ Todos los equipos medidos en los últimos 7 días.</p>"

    hoy_html = date.today()
    dias_desde_jue = (hoy_html.weekday() - 3) % 7
    jue_html = hoy_html - timedelta(days=dias_desde_jue)
    semana = jue_html.isocalendar()[1]
    return f"""<html><body style="font-family:Arial,sans-serif;background:#F1EFE8;color:#2C2C2A;padding:20px;">
    <div style="max-width:720px;margin:auto;">
      <div style="background:{TECK_GREEN};padding:20px;border-radius:12px;margin-bottom:20px;">
        <h1 style="color:white;margin:0;">GET Wear Monitor</h1>
        <p style="color:rgba(255,255,255,.9);margin:4px 0 0 0;">Reporte Semana {semana} · {date.today()}</p>
        <p style="color:rgba(255,255,255,.75);margin:2px 0 0 0;font-size:13px;">Teck QB2 · Pablo Cortés Ramos</p>
      </div>
      <h2 style="color:{TECK_GREEN};">Estado de flota</h2>
      <table style="width:100%;border-collapse:collapse;background:#FFFFFF;border-radius:8px;border:1px solid #D3D1C7;">
        <thead><tr style="background:{TECK_GREEN};color:white;">
          <th style="padding:10px;text-align:left;">Equipo</th><th style="padding:10px;text-align:left;">Última medición</th>
          <th style="padding:10px;text-align:left;">mm</th><th style="padding:10px;text-align:left;">Desgaste</th>
          <th style="padding:10px;text-align:left;">Estado</th><th style="padding:10px;text-align:left;">Días a crítico</th>
        </tr></thead><tbody>{filas}</tbody>
      </table>
      {tabla_sm}
      <p style="margin-top:30px;color:#888780;font-size:12px;">GET Wear Monitor · Teck QB2</p>
    </div></body></html>"""


def enviar_correo(destinatarios, html, excel_data, nombre_excel):
    try:
        smtp_user = st.secrets.get("SMTP_USER","")
        smtp_pass = st.secrets.get("SMTP_PASS","")
        smtp_host = st.secrets.get("SMTP_HOST","smtp.gmail.com")
        smtp_port = int(st.secrets.get("SMTP_PORT", 587))
        if not smtp_user or not smtp_pass:
            return False, "Configura SMTP_USER y SMTP_PASS en Secrets."
        semana = datetime.now().isocalendar()[1]
        msg = MIMEMultipart("mixed")
        msg["From"] = smtp_user
        msg["To"] = ", ".join(destinatarios)
        msg["Subject"] = f"GET Wear Monitor · Reporte Semana {semana} · {date.today()}"
        msg.attach(MIMEText(html, "html"))
        part = MIMEBase("application","octet-stream")
        part.set_payload(excel_data)
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{nombre_excel}"')
        msg.attach(part)
        with smtplib.SMTP(smtp_host, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_pass)
            server.sendmail(smtp_user, destinatarios, msg.as_string())
        return True, "Correo enviado correctamente."
    except Exception as e:
        return False, str(e)


# =========================================================
# REPORTE EJECUTIVO DOCX
# =========================================================
def generar_reporte_ejecutivo_docx(flota_data: list, periodo: str, fecha_str: str, semana: int, anio: int) -> bytes:
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    TECK_BLUE  = RGBColor(0x1F, 0x6F, 0xAE)
    TECK_GREEN_docx = RGBColor(0x00, 0x7A, 0x3D)
    TECK_DARK_docx  = RGBColor(0x0D, 0x1F, 0x2D)
    GRAY       = RGBColor(0x4A, 0x55, 0x68)

    OK_BG    = "EAF3DE"; OK_FG    = RGBColor(0x27, 0x50, 0x0A)
    MED_BG   = "FAEEDA"; MED_FG   = RGBColor(0x63, 0x38, 0x06)
    ALT_BG   = "FCEBEB"; ALT_FG   = RGBColor(0x71, 0x2B, 0x13)
    CRIT_BG  = "A32D2D"; CRIT_FG  = RGBColor(255, 255, 255)
    DIAS_OK  = "EAF3DE"; DIAS_MED = "FAEEDA"; DIAS_CRIT = "A32D2D"
    LIGHT_BG = "F4FAF0"

    def status_colors(estado):
        m = {"OK":(OK_BG,OK_FG),"MEDIO":(MED_BG,MED_FG),"ALTO":(ALT_BG,ALT_FG),"CRÍTICO":(CRIT_BG,CRIT_FG)}
        return m.get(estado,(OK_BG, OK_FG))

    def dias_colors(dias):
        if dias is None: return (OK_BG, OK_FG)
        d = float(dias)
        if d <= 10:  return (DIAS_OK,  RGBColor(0x27,0x50,0x0A))
        if d <= 14:  return (DIAS_MED, RGBColor(0x63,0x38,0x06))
        return (DIAS_CRIT, RGBColor(255,255,255))

    def pct_colors(pct):
        if pct is None: return (OK_BG, OK_FG)
        p = float(pct)
        if p >= 75: return (CRIT_BG, CRIT_FG)
        if p >= 45: return (MED_BG,  MED_FG)
        return (OK_BG, OK_FG)

    def set_cell_bg(cell, hex_color):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def add_cell(table_cell, text, bold=False, color=None, bg=None, size=9, align=WD_ALIGN_PARAGRAPH.LEFT, italics=False):
        if bg: set_cell_bg(table_cell, bg)
        table_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p_cell = table_cell.paragraphs[0]
        p_cell.alignment = align
        p_cell.paragraph_format.space_before = Pt(2)
        p_cell.paragraph_format.space_after = Pt(2)
        run = p_cell.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italics
        if color: run.font.color.rgb = color
        return run

    doc = DocxDoc()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("REPORTE EJECUTIVO")
    run.font.name = "Arial"; run.font.size = Pt(28); run.font.bold = True
    run.font.color.rgb = TECK_DARK_docx

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub_p.add_run("GET Wear Monitor  ·  Teck QB2")
    run2.font.name = "Arial"; run2.font.size = Pt(16)
    run2.font.color.rgb = TECK_BLUE

    per_p = doc.add_paragraph()
    per_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = per_p.add_run(f"Período: {periodo}  ·  Semana {semana}/{anio}  ·  {fecha_str}")
    run3.font.name = "Arial"; run3.font.size = Pt(11); run3.font.italic = True
    run3.font.color.rgb = GRAY

    doc.add_paragraph()

    kpi_ok   = len([f for f in flota_data if f["estado"]=="OK"])
    kpi_med  = len([f for f in flota_data if f["estado"]=="MEDIO"])
    kpi_alto = len([f for f in flota_data if f["estado"]=="ALTO"])
    kpi_crit = len([f for f in flota_data if f["estado"]=="CRÍTICO"])

    h = doc.add_heading("1. Resumen Ejecutivo", level=1)
    h.runs[0].font.color.rgb = TECK_BLUE

    kpi_table = doc.add_table(rows=2, cols=4)
    kpi_table.style = "Table Grid"
    kpi_data = [
        (str(kpi_ok),   "🟢 OK",      "Operación normal",   OK_BG,   OK_FG),
        (str(kpi_med),  "🟡 Monitoreo","Atención requerida", MED_BG,  MED_FG),
        (str(kpi_alto), "🟠 Programar","Coordinar cambio",   ALT_BG,  ALT_FG),
        (str(kpi_crit), "🔴 Crítico",  "Acción inmediata",   CRIT_BG, CRIT_FG),
    ]
    for i,(num,lbl,desc,bg,fg) in enumerate(kpi_data):
        c = kpi_table.cell(0,i); set_cell_bg(c, bg)
        p_c = c.paragraphs[0]; p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p_c.add_run(num); r.font.name="Arial"; r.font.size=Pt(28); r.font.bold=True; r.font.color.rgb=fg
        c2 = kpi_table.cell(1,i); set_cell_bg(c2, bg)
        p_c2 = c2.paragraphs[0]; p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p_c2.add_run(f"{lbl}\n{desc}"); r2.font.name="Arial"; r2.font.size=Pt(9); r2.font.bold=True; r2.font.color.rgb=fg

    doc.add_paragraph()

    h2 = doc.add_heading("2. Estado Completo de Flota", level=1)
    h2.runs[0].font.color.rgb = TECK_BLUE

    flota_sorted = sorted(flota_data, key=lambda x: x["pct"], reverse=True)
    t = doc.add_table(rows=1+len(flota_sorted), cols=7)
    t.style = "Table Grid"

    headers = ["Equipo","Tipo","mm actual","Desgaste %","Estado","Días a crítico","Prioridad"]
    widths_cm = [1.5, 3.5, 2.0, 2.0, 2.5, 2.5, 2.5]
    for i,(hdr,w) in enumerate(zip(headers,widths_cm)):
        cell_h = t.cell(0,i)
        set_cell_bg(cell_h, "007A3D")
        add_cell(cell_h, hdr, bold=True, color=RGBColor(255,255,255), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, f in enumerate(flota_sorted):
        row_bg = LIGHT_BG if ri%2==0 else "FFFFFF"
        bg_e, fg_e = status_colors(f["estado"])
        bg_p, fg_p = pct_colors(f["pct"])
        bg_d, fg_d = dias_colors(f["d_crit"])

        prio = "Normal"; prio_bg = OK_BG; prio_fg = OK_FG
        if f["estado"]=="CRÍTICO": prio="URGENTE"; prio_bg=CRIT_BG; prio_fg=CRIT_FG
        elif f["estado"]=="ALTO":  prio="ESTA SEMANA"; prio_bg=ALT_BG; prio_fg=ALT_FG
        elif f["d_crit"] and f["d_crit"]<=14: prio="PRÓX. SEMANA"; prio_bg=MED_BG; prio_fg=MED_FG

        vals = [
            (f["equipo"],                    row_bg, GRAY,   True, WD_ALIGN_PARAGRAPH.CENTER),
            (f["tipo"],                      row_bg, GRAY,   False,WD_ALIGN_PARAGRAPH.LEFT),
            (f"{f['mm']:.0f} mm",            row_bg, GRAY,   False,WD_ALIGN_PARAGRAPH.CENTER),
            (f"{f['pct']:.1f}%",             bg_p,   fg_p,   True, WD_ALIGN_PARAGRAPH.CENTER),
            (f["estado"],                    bg_e,   fg_e,   True, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{f['d_crit']:.0f} días" if f["d_crit"] is not None else "—", bg_d, fg_d, True, WD_ALIGN_PARAGRAPH.CENTER),
            (prio,                           prio_bg,prio_fg,True, WD_ALIGN_PARAGRAPH.CENTER),
        ]
        for ci,(txt,bg,fg,bld,aln) in enumerate(vals):
            add_cell(t.cell(ri+1,ci), txt, bold=bld, color=fg, bg=bg, size=9, align=aln)

    doc.add_paragraph()

    h3 = doc.add_heading("3. Proyección de Cambios (próximos 60 días)", level=1)
    h3.runs[0].font.color.rgb = TECK_BLUE

    flota_proj = sorted([f for f in flota_data if f["d_crit"] is not None], key=lambda x: x["d_crit"])
    if flota_proj:
        tp = doc.add_table(rows=1+len(flota_proj), cols=5)
        tp.style = "Table Grid"
        for i,hdr in enumerate(["Equipo","Tipo","Días a crítico","Estado","Prioridad"]):
            ch = tp.cell(0,i); set_cell_bg(ch,"007A3D")
            add_cell(ch, hdr, bold=True, color=RGBColor(255,255,255), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        for ri,f in enumerate(flota_proj):
            bg_e, fg_e = status_colors(f["estado"])
            bg_d, fg_d = dias_colors(f["d_crit"])
            prio="Normal"; prio_bg=OK_BG; prio_fg=OK_FG
            if f["estado"]=="CRÍTICO": prio="🔴 URGENTE"; prio_bg=CRIT_BG; prio_fg=CRIT_FG
            elif f["estado"]=="ALTO":  prio="🟠 ESTA SEMANA"; prio_bg=ALT_BG; prio_fg=ALT_FG
            elif f["d_crit"]<=14:      prio="🟡 PRÓX. SEMANA"; prio_bg=MED_BG; prio_fg=MED_FG
            row_bg = LIGHT_BG if ri%2==0 else "FFFFFF"
            vals2 = [
                (f["equipo"],row_bg,GRAY,True,WD_ALIGN_PARAGRAPH.CENTER),
                (f["tipo"],row_bg,GRAY,False,WD_ALIGN_PARAGRAPH.LEFT),
                (f"{f['d_crit']:.0f} días",bg_d,fg_d,True,WD_ALIGN_PARAGRAPH.CENTER),
                (f["estado"],bg_e,fg_e,True,WD_ALIGN_PARAGRAPH.CENTER),
                (prio,prio_bg,prio_fg,True,WD_ALIGN_PARAGRAPH.CENTER),
            ]
            for ci,(txt,bg,fg,bld,aln) in enumerate(vals2):
                add_cell(tp.cell(ri+1,ci), txt, bold=bld, color=fg, bg=bg, size=9, align=aln)

    doc.add_paragraph()

    h4 = doc.add_heading("4. Conclusiones y Recomendaciones", level=1)
    h4.runs[0].font.color.rgb = TECK_BLUE

    criticos_list = [f for f in flota_data if f["estado"] in ("CRÍTICO","ALTO")]
    medio_list    = [f for f in flota_data if f["estado"]=="MEDIO"]

    if criticos_list:
        p_c = doc.add_paragraph()
        r_c = p_c.add_run("🔴  ACCIÓN INMEDIATA: ")
        r_c.font.name="Arial"; r_c.font.bold=True; r_c.font.color.rgb=CRIT_FG; r_c.font.size=Pt(10)
        r_c2 = p_c.add_run(", ".join([f"Equipo {f['equipo']} ({f['estado']})" for f in criticos_list]))
        r_c2.font.name="Arial"; r_c2.font.size=Pt(10); r_c2.font.color.rgb=CRIT_FG

    if medio_list:
        p_m = doc.add_paragraph()
        r_m = p_m.add_run("🟡  MONITOREO: ")
        r_m.font.name="Arial"; r_m.font.bold=True; r_m.font.color.rgb=MED_FG; r_m.font.size=Pt(10)
        r_m2 = p_m.add_run(", ".join([f"Equipo {f['equipo']} ({f['pct']:.0f}% desgaste)" for f in medio_list]))
        r_m2.font.name="Arial"; r_m2.font.size=Pt(10); r_m2.font.color.rgb=MED_FG

    p_ok = doc.add_paragraph()
    r_ok = p_ok.add_run(f"🟢  FLOTA OK: {kpi_ok} equipos en condición normal. Continuar monitoreo semanal.")
    r_ok.font.name="Arial"; r_ok.font.size=Pt(10); r_ok.font.color.rgb=OK_FG

    doc.add_paragraph()
    p_firma = doc.add_paragraph()
    p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_firma = p_firma.add_run(f"Pablo Cortés Ramos  ·  Ingeniero de Mantenimiento / Confiabilidad  ·  Teck QB2  ·  {fecha_str}")
    r_firma.font.name="Arial"; r_firma.font.size=Pt(9); r_firma.font.color.rgb=GRAY; r_firma.font.italic=True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =========================================================
# PROYECCIÓN DE FECHAS DE CAMBIO
# =========================================================
MIN_MEDS_PROYECCION = 2

def proyectar_fecha_cambio(equipo: str) -> dict:
    try:
        uc = ultimo_cambio_equipo(equipo)
        q = sb().table("mediciones").select("fecha,horometro,mm_usada,sospechoso,comentario_sospecha").eq("equipo", equipo).eq("es_cambio", False).order("fecha", desc=False)
        if uc:
            q = q.gte("fecha", uc["fecha"])
        todas = q.execute().data or []

        sospechosas = [m for m in todas if m.get("sospechoso")]
        meds = [m for m in todas if not m.get("sospechoso")]

        regla      = regla_por_equipo(equipo)
        cfg        = REGLAS[regla]
        mm_critico = cfg["mm_critico"]
        tasa_default = 0.028 if regla == "MOTONIVELADORA" else 0.013

        df_horo = cargar_horometros_db()
        h_dia = None; horo_excel = None
        if not df_horo.empty:
            row_h = df_horo[df_horo["equipo"] == equipo]
            if not row_h.empty:
                horo_excel = float(row_h.iloc[0]["horometro_actual"])
                prom = row_h.iloc[0].get("promedio_30d") or row_h.iloc[0].get("promedio_historico")
                if prom and float(prom) > 0:
                    h_dia = float(prom)

        nota_sospechosas = f"{len(sospechosas)} medición(es) sospechosa(s) excluida(s) del cálculo." if sospechosas else ""

        if len(meds) < 1:
            return {"ok": False, "error": "Sin mediciones válidas en el ciclo actual." + (f" ({nota_sospechosas})" if nota_sospechosas else ""), "meds": 0}

        if len(meds) == 1 and h_dia:
            mm_actual = float(meds[0]["mm_usada"])
            restante_mm = mm_actual - mm_critico
            if restante_mm <= 0:
                return {"ok": True, "estado": "CRÍTICO", "mm_actual": mm_actual, "mm_critico": mm_critico,
                        "tasa": None, "tasa_fuente": "Sin tasa", "horas_restantes": 0, "dias_restantes": 0,
                        "fecha_cambio": date.today(), "meds": len(meds), "h_dia": h_dia,
                        "confianza": "Muy baja — solo 1 medición válida", "confianza_nivel": 0,
                        "nota": nota_sospechosas, "anomalias": [], "fecha_ultima_med": meds[0]["fecha"]}
            horas_restantes = restante_mm / tasa_default
            dias_restantes = horas_restantes / h_dia
            return {"ok": True,
                    "estado": "OK" if dias_restantes > 30 else ("MEDIO" if dias_restantes > 14 else ("ALTO" if dias_restantes > 7 else "CRÍTICO")),
                    "mm_actual": mm_actual, "mm_critico": mm_critico,
                    "tasa": round(tasa_default, 5), "tasa_fuente": "Tasa por defecto (1 medición válida)",
                    "horas_restantes": round(horas_restantes, 1), "dias_restantes": round(dias_restantes, 1),
                    "fecha_cambio": date.today() + timedelta(days=int(dias_restantes)),
                    "meds": len(meds), "h_dia": round(h_dia, 1),
                    "confianza": "Muy baja — 1 medición válida, tasa estimada", "confianza_nivel": 1,
                    "nota": nota_sospechosas, "anomalias": [], "fecha_ultima_med": meds[0]["fecha"]}

        if len(meds) < 2:
            return {"ok": False, "error": f"Solo {len(meds)} medición válida y sin horómetros Excel para estimar." + (f" {nota_sospechosas}" if nota_sospechosas else ""), "meds": len(meds)}

        tasas_validas = []; anomalias = []
        for i in range(len(meds) - 1):
            dh  = float(meds[i+1]["horometro"]) - float(meds[i]["horometro"])
            dmm = float(meds[i]["mm_usada"])     - float(meds[i+1]["mm_usada"])
            if dh <= 0:
                anomalias.append(f"Horómetro no sube entre medición {i+1} y {i+2}")
                continue
            if dmm < 0:
                anomalias.append(f"mm sube {-dmm:.1f} mm entre medición {i+1} y {i+2} — posible cambio no registrado")
                continue
            if dmm == 0:
                continue
            tasa_i = dmm / dh
            if tasa_i > tasa_default * 15:
                anomalias.append(f"Tasa anómala {tasa_i:.4f} mm/h en intervalo {i+1}-{i+2}")
                continue
            tasas_validas.append(tasa_i)

        mm_actual    = float(meds[-1]["mm_usada"])
        horo_actual  = horo_excel or float(meds[-1]["horometro"])
        fecha_ultima = meds[-1]["fecha"]
        restante_mm  = mm_actual - mm_critico

        if not tasas_validas:
            tasa = tasa_default; tasa_fuente = "Tasa por defecto (datos anómalos)"; confianza = "Muy baja — revisar mediciones"; confianza_nivel = 1
        else:
            tasa = sum(tasas_validas) / len(tasas_validas)
            n_valid = len(tasas_validas); n_total = len(meds) - 1
            tasa_fuente = f"Tasa real ({n_valid}/{n_total} intervalos válidos)" if n_valid < n_total else "Tasa real (todas las mediciones)"
            if len(meds) >= 5 and n_valid >= 4 and not anomalias:
                confianza = "Alta"; confianza_nivel = 4
            elif len(meds) >= 3 and n_valid >= 2:
                confianza = "Media" + (" — hay anomalías" if anomalias else ""); confianza_nivel = 3
            else:
                confianza = "Baja — pocas mediciones"; confianza_nivel = 2

        if restante_mm <= 0:
            return {"ok": True, "estado": "CRÍTICO", "mm_actual": mm_actual, "mm_critico": mm_critico,
                    "tasa": round(tasa, 5), "tasa_fuente": tasa_fuente, "horas_restantes": 0, "dias_restantes": 0,
                    "fecha_cambio": date.today(), "meds": len(meds), "h_dia": h_dia,
                    "confianza": confianza, "confianza_nivel": confianza_nivel,
                    "anomalias": anomalias, "nota": nota_sospechosas, "fecha_ultima_med": fecha_ultima}

        horas_restantes = restante_mm / tasa
        dias_restantes  = horas_restantes / h_dia if h_dia and h_dia > 0 else horas_restantes / 12
        fecha_cambio    = date.today() + timedelta(days=int(dias_restantes))

        return {"ok": True,
                "estado": "OK" if dias_restantes > 30 else ("MEDIO" if dias_restantes > 14 else ("ALTO" if dias_restantes > 7 else "CRÍTICO")),
                "mm_actual": mm_actual, "mm_critico": mm_critico,
                "tasa": round(tasa, 5), "tasa_fuente": tasa_fuente,
                "horas_restantes": round(horas_restantes, 1), "dias_restantes": round(dias_restantes, 1),
                "fecha_cambio": fecha_cambio, "meds": len(meds),
                "h_dia": round(h_dia, 1) if h_dia else None,
                "confianza": confianza, "confianza_nivel": confianza_nivel,
                "anomalias": anomalias, "nota": nota_sospechosas,
                "horo_actual": horo_actual, "fecha_ultima_med": fecha_ultima}

    except Exception as e:
        return {"ok": False, "error": str(e), "meds": 0}


# =========================================================
# APP UI
# =========================================================
ADMIN_PASSWORD = st.secrets.get("ADMIN_PASSWORD", "254828")

with st.sidebar:
    if st.button("🔄 Actualizar app", help="Recarga completa — elimina errores visuales tras actualizaciones"):
        st.cache_data.clear()
        st.cache_resource.clear()
        st.session_state.pop("app_version", None)
        st.query_params["v"] = APP_VERSION
        st.rerun()

    st.divider()
    st.subheader("Administración")
    admin_ok = False
    pwd = st.text_input("Clave administrador", type="password")
    if pwd:
        admin_ok = (pwd == ADMIN_PASSWORD)
        if admin_ok:
            st.success("✅ Modo administrador activo")
        else:
            st.error("Clave incorrecta")


tabs_base = [
    "📏 Ingreso Medición",
    "🔄 Cambio de Cuchilla",
    "📋 Historial",
    "🚛 Estado Flota",
    "📂 Horómetros Excel",
    "📊 Reporte Semanal",
    "📄 Reporte Ejecutivo",
]
if admin_ok:
    tabs_base.append("📆 Proyección Cambios")
    tabs_base.append("👥 Accesos")

tabs = st.tabs(tabs_base)
TAB_PROYECCION = 7 if admin_ok else None
TAB_ACCESOS    = 8 if admin_ok else None


# ─────────────────────────────────────────────
# TAB 1: INGRESO MEDICIÓN
# ─────────────────────────────────────────────
with tabs[0]:
    c1, c2 = st.columns([1,1], gap="large")
    with c1:
        st.subheader("Ingreso de medición")
        equipo    = st.selectbox("Equipo", EQUIPOS, key="eq_m")
        fecha_m   = st.date_input("Fecha", value=date.today(), key="f_m")
        usuario_m = st.text_input("Técnico", key="u_m")
        st.divider()
        regla   = regla_por_equipo(equipo)
        mm_max  = REGLAS[regla]["mm_nuevo"]

        df_horo = cargar_horometros_db()
        h_sugerido = None
        if not df_horo.empty:
            row_h = df_horo[df_horo["equipo"] == equipo]
            if not row_h.empty:
                h_sugerido = float(row_h.iloc[0]["horometro_actual"])
                fecha_excel = row_h.iloc[0]["fecha"]
                st.caption(f"📊 Excel horómetros: **{h_sugerido:,.0f} hrs** al {fecha_excel}")

        horometro_m = st.number_input("Horómetro", min_value=0.0, step=1.0, key="h_m",
                                       value=float(h_sugerido) if h_sugerido else 0.0)

        if horometro_m > 0 and h_sugerido:
            val = validar_horometro(equipo, horometro_m)
            if not val["ok"]:
                st.warning(
                    f"⚠️ **Horómetro inusual** — diferencia de {val['diferencia']:,.0f} hrs "
                    f"({val['diferencia_pct']}%) respecto al Excel ({val['h_excel']:,.0f} hrs al {val['fecha_excel']}). "
                    f"Verifica antes de guardar."
                )

        mm_izq_m = st.number_input("Medición IZQ (mm)", min_value=0.0, value=mm_max, step=0.1, key="mi_m")
        mm_der_m = st.number_input("Medición DER (mm)", min_value=0.0, value=mm_max, step=0.1, key="md_m")
        st.caption("Se usa el valor MENOR (más crítico) para evaluar.")

        uc = ultimo_cambio_equipo(equipo)
        if uc:
            st.info(f"📌 Último cambio: {uc['fecha']} · Horómetro: {uc['horometro']:,.0f} hrs")

        mm_min_ingreso = min(mm_izq_m, mm_der_m)
        sospecha_info  = verificar_sospecha(equipo, mm_min_ingreso)

        if sospecha_info["sospechoso"] and "confirmar_sospecha" not in st.session_state:
            st.warning(
                f"⚠️ **¿Estás seguro?** El valor subió **{sospecha_info['diferencia']:.1f} mm** "
                f"({sospecha_info['mm_anterior']:.1f} → {sospecha_info['mm_nueva']:.1f} mm) "
                f"sin un cambio de cuchilla registrado. Esto es inusual."
            )
            col_si, col_no = st.columns(2)
            with col_si:
                if st.button("✅ Sí, el valor es correcto — guardar igual", key="btn_confirmar_sospecha"):
                    st.session_state["confirmar_sospecha"] = True
                    st.session_state["guardar_pendiente"]  = True
                    st.rerun()
            with col_no:
                if st.button("✏️ No, quiero corregir el valor", key="btn_corregir_sospecha"):
                    st.session_state.pop("confirmar_sospecha", None)
                    st.session_state.pop("guardar_pendiente", None)
                    st.info("Corrige los valores de medición arriba.")
        else:
            if st.button("Guardar medición", type="primary", key="btn_m"):
                st.session_state["guardar_pendiente"] = True

        if st.session_state.pop("guardar_pendiente", False):
            errores = []
            if not usuario_m.strip(): errores.append("Ingresa el nombre del técnico.")
            if horometro_m <= 0:      errores.append("Ingresa un horómetro válido.")
            mn, mx = rango_regla(regla)
            if not (mn <= mm_izq_m <= mx): errores.append(f"IZQ fuera de rango ({mn}–{mx} mm).")
            if not (mn <= mm_der_m <= mx): errores.append(f"DER fuera de rango ({mn}–{mx} mm).")
            if errores:
                for e in errores: st.error(e)
                st.session_state.pop("confirmar_sospecha", None)
            else:
                try:
                    r = evaluar(equipo, horometro_m, mm_izq_m, mm_der_m)
                    es_sospechosa  = sospecha_info["sospechoso"] and st.session_state.pop("confirmar_sospecha", False)
                    comentario_s   = sospecha_info.get("comentario_auto", "") if es_sospechosa else ""
                    guardar_medicion(fecha_m, equipo, horometro_m, mm_izq_m, mm_der_m, usuario_m, r,
                                     sospechoso=es_sospechosa, comentario_sospecha=comentario_s)
                    if es_sospechosa:
                        st.warning(f"⚠️ Medición guardada y marcada como **sospechosa**: {comentario_s}")
                    else:
                        st.success("✅ Medición guardada.")
                    ca, cb, cc = st.columns(3)
                    ca.metric("Estado", f"{COLOR_ESTADO.get(r['estado'],'⚪')} {r['estado']}")
                    cb.metric("Desgaste %", f"{r['pct']:.1f}%")
                    cc.metric("mm usada", f"{r['mm_usada']:.1f}")
                    if r.get("tasa"):    st.info(f"Tasa: **{r['tasa']} mm/h** · {r.get('fuente_tasa','')}")
                    if r.get("h_critico") is not None:
                        st.warning(f"⏱ Proyección a crítico: ~{r['h_critico']} h / ~{r['d_critico']} días")
                except Exception as e:
                    st.error(f"Error: {e}")

    with c2:
        st.subheader(f"Curva de desgaste — Equipo {equipo}")
        df_curva = cargar_historial(500)
        df_eq = pd.DataFrame()
        if not df_curva.empty and "equipo" in df_curva.columns:
            df_eq = df_curva[df_curva["equipo"] == equipo].copy()
            if "es_cambio" in df_eq.columns:
                df_eq = df_eq[df_eq["es_cambio"] == False]
        if not df_eq.empty and "horometro" in df_eq.columns and "mm_usada" in df_eq.columns:
            df_eq = df_eq.dropna(subset=["horometro","mm_usada"]).sort_values("horometro")
            cfg_eq = REGLAS[regla_por_equipo(equipo)]

            try:
                import plotly.graph_objects as go
                df_plot = df_eq[["horometro","mm_usada","fecha"]].copy()
                df_plot["sospechoso"] = df_eq["sospechoso"].fillna(False).astype(bool) if "sospechoso" in df_eq.columns else False
                df_plot["comentario"] = df_eq["comentario_sospecha"].fillna("") if "comentario_sospecha" in df_eq.columns else ""

                df_ok  = df_plot[~df_plot["sospechoso"]]
                df_mal = df_plot[df_plot["sospechoso"]]

                fig = go.Figure()
                fig.add_trace(go.Scatter(
                    x=df_plot["horometro"], y=df_plot["mm_usada"],
                    mode="lines", line=dict(color=TECK_GREEN, width=2),
                    name="Desgaste", showlegend=False,
                ))
                fig.add_trace(go.Scatter(
                    x=df_ok["horometro"], y=df_ok["mm_usada"],
                    mode="markers",
                    marker=dict(color=TECK_GREEN_2, size=8, line=dict(color=TECK_GREEN, width=1)),
                    name="Medición OK",
                    customdata=df_ok[["fecha","comentario"]].values,
                    hovertemplate="<b>Horómetro:</b> %{x:,.0f} hrs<br><b>mm:</b> %{y:.1f}<br><b>Fecha:</b> %{customdata[0]}<extra></extra>",
                ))
                if not df_mal.empty:
                    fig.add_trace(go.Scatter(
                        x=df_mal["horometro"], y=df_mal["mm_usada"],
                        mode="markers",
                        marker=dict(color="#A32D2D", size=14, symbol="x", line=dict(color="#A32D2D", width=3)),
                        name="⚠️ Sospechoso",
                        customdata=df_mal[["fecha","comentario"]].values,
                        hovertemplate="<b>⚠️ SOSPECHOSO</b><br>Horómetro: %{x:,.0f}<br>mm: %{y:.1f}<br>Fecha: %{customdata[0]}<br>%{customdata[1]}<extra></extra>",
                    ))
                fig.add_hline(y=cfg_eq["mm_critico"], line_dash="dash", line_color="#A32D2D",
                    annotation_text=f"Límite crítico {cfg_eq['mm_critico']} mm",
                    annotation_position="bottom right", annotation_font_color="#A32D2D")
                fig.add_hline(y=cfg_eq["mm_nuevo"], line_dash="dot", line_color="#3B6D11",
                    annotation_text=f"GET nuevo {cfg_eq['mm_nuevo']} mm",
                    annotation_position="top right", annotation_font_color="#3B6D11")
                fig.update_layout(
                    height=300, margin=dict(l=0, r=0, t=10, b=0),
                    paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(255,255,255,0.8)",
                    font=dict(color="#2C2C2A"),
                    xaxis=dict(title="Horómetro (hrs)", gridcolor="#D3D1C7", color="#444441"),
                    yaxis=dict(title="mm usada", gridcolor="#D3D1C7", color="#444441"),
                    legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1),
                )
                st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
            except Exception as ex:
                st.line_chart(df_eq.set_index("horometro")[["mm_usada"]], use_container_width=True)
                st.caption(f"(modo simple — {ex})")

            st.caption(f"🔴 Límite crítico: {cfg_eq['mm_critico']} mm · 🟢 GET nuevo: {cfg_eq['mm_nuevo']} mm · ✖ Punto sospechoso (excluido de proyección)")
        else:
            st.info("Sin datos suficientes para la curva.")

        st.subheader("Últimas mediciones")
        if not df_eq.empty:
            cols_s = [c for c in ["fecha","horometro","mm_izq","mm_der","mm_usada","condicion_pct","estado","sospechoso","comentario_sospecha","usuario"] if c in df_eq.columns]
            df_show = df_eq[cols_s].sort_values("fecha", ascending=False).head(10).copy()
            if "sospechoso" in df_show.columns:
                df_show["⚠️"] = df_show["sospechoso"].apply(lambda x: "🔴 Sospechoso" if x else "")
                cols_display = ["fecha","horometro","mm_izq","mm_der","mm_usada","condicion_pct","estado","⚠️","comentario_sospecha","usuario"]
                cols_display = [c for c in cols_display if c in df_show.columns]
                df_show = df_show[cols_display]

            def style_sospechosa(row):
                if row.get("⚠️","") == "🔴 Sospechoso":
                    return [f"background-color:{BG_CRIT};color:{TEXT_CRIT}"] * len(row)
                return [""] * len(row)

            st.dataframe(
                df_show.style.apply(style_sospechosa, axis=1),
                use_container_width=True
            )
        else:
            st.info(f"Sin mediciones para equipo {equipo}.")


# ─────────────────────────────────────────────
# TAB 2: CAMBIO DE CUCHILLA
# ─────────────────────────────────────────────
with tabs[1]:
    sub1, sub2 = st.tabs(["➕ Registrar cambio", "📋 Historial de cambios"])

    with sub1:
        st.subheader("Registrar cambio de cuchilla / GET nuevo")
        c1, c2 = st.columns(2, gap="large")
        with c1:
            eq_c     = st.selectbox("Equipo", EQUIPOS, key="eq_c")
            fecha_c  = st.date_input("Fecha de cambio", value=date.today(), key="f_c")
            hr_c     = st.number_input("Horómetro al cambio", min_value=0.0, step=1.0, key="h_c")
            regla_c  = regla_por_equipo(eq_c)
            mm_c     = REGLAS[regla_c]["mm_critico"]
            mm_max_c = REGLAS[regla_c]["mm_nuevo"]
            mm_izq_c = st.number_input("IZQ final GET retirado (mm)", min_value=0.0, max_value=float(mm_max_c), value=float(mm_c), step=0.1, key="mi_c")
            mm_der_c = st.number_input("DER final GET retirado (mm)", min_value=0.0, max_value=float(mm_max_c), value=float(mm_c), step=0.1, key="md_c")
            virada_c = st.radio("¿Fue virada?", ["NO","SÍ"], horizontal=True, key="v_c")
            motivo_c = st.selectbox("Motivo", ["Desgaste normal","Preventivo","Daño / impacto","Campaña mantenimiento","Otro"], key="mot_c")
        with c2:
            tec1_c = st.text_input("Técnico 1", key="t1_c")
            tec2_c = st.text_input("Técnico 2 (opcional)", key="t2_c")
            obs_c  = st.text_area("Observaciones / OT generadas", height=120, key="obs_c")
            sup_c  = st.text_input("Registrado por", key="sup_c")
            st.divider()
            st.info(f"Al confirmar: ciclo nuevo con **{mm_max_c:.0f} mm**")
            if st.button("✅ Confirmar cambio", type="primary", key="btn_c"):
                errores = []
                if hr_c <= 0:          errores.append("Horómetro debe ser > 0.")
                if not tec1_c.strip(): errores.append("Ingresa al menos el Técnico 1.")
                if not sup_c.strip():  errores.append("Ingresa quién registra.")
                if errores:
                    for e in errores: st.error(e)
                else:
                    try:
                        guardar_cambio(fecha_c, eq_c, hr_c, mm_izq_c, mm_der_c,
                                       virada_c=="SÍ", motivo_c, obs_c, tec1_c, tec2_c, sup_c)
                        st.success(f"🔄 Cambio registrado · Equipo **{eq_c}** · Horómetro **{hr_c:,.0f} hrs**")
                        st.balloons()
                    except Exception as e:
                        st.error(f"Error: {e}")

    with sub2:
        st.subheader("Historial de cambios")
        df_cam = cargar_cambios()
        if not df_cam.empty:
            df_cam["fue_virada"] = df_cam["fue_virada"].map({True:"SÍ",False:"NO",1:"SÍ",0:"NO"})
            cols_c = [c for c in ["fecha","equipo","horometro","mm_izq_final","mm_der_final","fue_virada","motivo","tecnico_1","tecnico_2","observaciones","usuario"] if c in df_cam.columns]
            st.dataframe(df_cam[cols_c], use_container_width=True)
            st.download_button("⬇️ Descargar Excel", data=excel_bytes(df_cam[cols_c]), file_name="cambios.xlsx")
        else:
            st.info("Sin cambios registrados aún.")


# ─────────────────────────────────────────────
# TAB 3: HISTORIAL + EDICIÓN ADMIN
# ─────────────────────────────────────────────
with tabs[2]:
    st.subheader("📋 Historial completo")
    df_h = cargar_historial()
    if not df_h.empty:
        cf1, cf2 = st.columns(2)
        with cf1: eq_fil  = st.multiselect("Equipo",  sorted(df_h["equipo"].unique()), key="fil_eq")
        with cf2: est_fil = st.multiselect("Estado", sorted(df_h["estado"].dropna().unique()), key="fil_est")
        df_sh = df_h.copy()
        if eq_fil:  df_sh = df_sh[df_sh["equipo"].isin(eq_fil)]
        if est_fil: df_sh = df_sh[df_sh["estado"].isin(est_fil)]
        cols_h = [c for c in ["fecha","equipo","horometro","mm_izq","mm_der","mm_usada","condicion_pct","estado","usuario"] if c in df_sh.columns]
        st.dataframe(df_sh[cols_h], use_container_width=True)
        st.download_button("⬇️ Descargar Excel", data=excel_bytes(df_sh[cols_h]), file_name="historial.xlsx")

        if admin_ok:
            st.divider()
            st.markdown("### 🛠️ Gestión de mediciones — Administrador")
            st.caption("Selecciona una medición para editarla o eliminarla. Solo visible con clave de administrador.")

            df_edit = df_h[["id","fecha","equipo","horometro","mm_izq","mm_der","mm_usada","estado","usuario"]].copy()
            if "sospechoso" in df_h.columns:
                df_edit["sospechoso"] = df_h["sospechoso"].fillna(False)
            df_edit["descripcion"] = df_edit.apply(
                lambda r: f"ID {int(r['id'])} | {r['fecha']} | Equipo {r['equipo']} | Horo {r['horometro']:,.0f} hrs | {r['mm_usada']:.1f} mm | {r['usuario']}"
                          + (" 🔴" if r.get("sospechoso") else ""), axis=1
            )
            sel = st.selectbox("🔍 Seleccionar medición", ["— Selecciona una medición —"] + df_edit["descripcion"].tolist(), key="sel_edit")

            if sel and sel != "— Selecciona una medición —":
                id_edit  = int(sel.split("|")[0].replace("ID","").strip())
                row_edit = df_h[df_h["id"] == id_edit].iloc[0]
                eq_edit  = str(row_edit["equipo"])
                es_sosp  = bool(row_edit.get("sospechoso", False))
                com_sosp = str(row_edit.get("comentario_sospecha","") or "")

                bg_card, border_card, text_card = _ESTADO_CARD.get(str(row_edit.get("estado","")), (BG_NODAT, "#D3D1C7", "#444441"))
                st.markdown(f"""
                <div style="background:{bg_card};border:1px solid {border_card};border-radius:12px;padding:14px 18px;margin:10px 0 18px 0;">
                    <div style="display:flex;gap:24px;flex-wrap:wrap;align-items:center;">
                        <div><span style="font-size:11px;color:{text_card};opacity:.7;display:block;">EQUIPO</span><b style="font-size:20px;color:{text_card};">#{eq_edit}</b></div>
                        <div><span style="font-size:11px;color:{text_card};opacity:.7;display:block;">FECHA</span><b style="color:{text_card};">{row_edit['fecha']}</b></div>
                        <div><span style="font-size:11px;color:{text_card};opacity:.7;display:block;">HORÓMETRO</span><b style="color:{text_card};">{float(row_edit['horometro']):,.0f} hrs</b></div>
                        <div><span style="font-size:11px;color:{text_card};opacity:.7;display:block;">IZQ / DER</span><b style="color:{text_card};">{float(row_edit['mm_izq']):.1f} / {float(row_edit['mm_der']):.1f} mm</b></div>
                        <div><span style="font-size:11px;color:{text_card};opacity:.7;display:block;">ESTADO</span><b style="color:{text_card};">{row_edit.get('estado','—')}</b></div>
                        <div><span style="font-size:11px;color:{text_card};opacity:.7;display:block;">TÉCNICO</span><b style="color:{text_card};">{row_edit['usuario']}</b></div>
                    </div>
                    {"<div style=\"margin-top:10px;padding:8px 12px;background:" + BG_CRIT + ";border-radius:8px;font-size:13px;color:" + TEXT_CRIT + ";\">🔴 <b>Sospechosa:</b> " + com_sosp + "</div>" if es_sosp else ""}
                </div>
                """, unsafe_allow_html=True)

                tab_editar, tab_eliminar = st.tabs(["✏️ Editar medición", "🗑️ Eliminar medición"])

                with tab_editar:
                    ce1, ce2, ce3 = st.columns(3)
                    with ce1:
                        horo_edit = st.number_input("Horómetro corregido", value=float(row_edit["horometro"]), step=1.0, key="horo_edit")
                    with ce2:
                        mm_izq_edit = st.number_input("IZQ corregida (mm)", value=float(row_edit["mm_izq"]), step=0.1, key="mi_edit")
                    with ce3:
                        mm_der_edit = st.number_input("DER corregida (mm)", value=float(row_edit["mm_der"]), step=0.1, key="md_edit")

                    usuario_edit = st.text_input("Registrado por (corrección)", value=str(row_edit["usuario"]), key="u_edit")

                    if horo_edit > 0:
                        val_edit = validar_horometro(eq_edit, horo_edit)
                        if not val_edit["ok"]:
                            st.warning(f"⚠️ Horómetro difiere del Excel ({val_edit['h_excel']:,.0f} hrs). Verifica antes de guardar.")

                    col_g1, col_g2 = st.columns(2)
                    with col_g1:
                        if st.button("💾 Guardar corrección", type="primary", key="btn_edit"):
                            try:
                                actualizar_medicion(id_edit, horo_edit, mm_izq_edit, mm_der_edit, usuario_edit, eq_edit)
                                st.success(f"✅ Medición ID {id_edit} actualizada correctamente.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error: {e}")
                    with col_g2:
                        if es_sosp:
                            if st.button("✅ Quitar marca sospechosa", key="btn_no_sosp"):
                                marcar_sospechosa(id_edit, False)
                                st.success("Marca sospechosa eliminada.")
                                st.rerun()
                        else:
                            com_manual = st.text_input("Motivo (obligatorio para marcar sospechosa)", key="com_sosp_manual")
                            if st.button("🔴 Marcar como sospechosa", key="btn_sosp"):
                                if not com_manual.strip():
                                    st.error("Ingresa el motivo antes de marcar.")
                                else:
                                    marcar_sospechosa(id_edit, True, com_manual.strip())
                                    st.warning(f"Medición ID {id_edit} marcada como sospechosa.")
                                    st.rerun()

                with tab_eliminar:
                    st.error(f"⚠️ Estás a punto de eliminar permanentemente la medición **ID {id_edit}** del equipo **{eq_edit}** ({row_edit['fecha']}). Esta acción **no se puede deshacer**.")
                    st.markdown("Para confirmar, escribe **ELIMINAR** en el campo de abajo:")
                    confirm_text = st.text_input("Escribe ELIMINAR para confirmar", key="confirm_del_text", placeholder="ELIMINAR")
                    if confirm_text.strip().upper() == "ELIMINAR":
                        if st.button("🗑️ Eliminar definitivamente", type="primary", key="btn_del"):
                            try:
                                eliminar_medicion(id_edit)
                                st.success(f"✅ Medición ID {id_edit} eliminada correctamente.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Error al eliminar: {e}")
    else:
        st.info("Sin mediciones aún.")


# ─────────────────────────────────────────────
# TAB 4: ESTADO FLOTA
# ─────────────────────────────────────────────
with tabs[3]:
    st.subheader("🚛 Estado de flota")
    df_flota = cargar_historial(5000)
    ultimos  = ultimos_por_equipo(df_flota)

    if not ultimos.empty:
        ec = ultimos["estado"].value_counts()
        k1,k2,k3,k4 = st.columns(4)
        k1.metric("🟢 OK",               int(ec.get("OK",0)))
        k2.metric("🟡 Monitoreo",         int(ec.get("MEDIO",0)))
        k3.metric("🟠 Programar cambio",  int(ec.get("ALTO",0)))
        k4.metric("🔴 Crítico",           int(ec.get("CRÍTICO",0)))

        st.divider()
        st.subheader("Desgaste por equipo (%)")
        if "condicion_pct" in ultimos.columns:
            ranking = ultimos[["equipo","condicion_pct"]].sort_values("condicion_pct",ascending=False).set_index("equipo")
            st.bar_chart(ranking)

        st.divider()
        st.subheader("Días sin medición")
        df_dias = dias_sin_medicion(df_flota)

        def color_fila_dias(val):
            try:
                v = float(val)
                if v <= 10:  return f"color:{DIAS_OK_FG};font-weight:bold"
                if v <= 14:  return f"color:{DIAS_MD_FG};font-weight:bold"
                return f"color:{DIAS_CR_BG};font-weight:bold"
            except:
                return "color:#888780"

        st.dataframe(
            df_dias.style.map(color_fila_dias, subset=["dias_sin_medir"]),
            use_container_width=True
        )

        st.divider()
        st.subheader("Proyección de cambios")
        cols_p = [c for c in ["equipo","mm_usada","condicion_pct","estado","tasa_mm_h","horas_a_critico","dias_a_critico"] if c in ultimos.columns]
        proy = ultimos[cols_p].sort_values("horas_a_critico", na_position="last") if "horas_a_critico" in cols_p else ultimos[cols_p]
        st.dataframe(proy, use_container_width=True)
    else:
        st.info("Sin datos de flota aún.")


# ─────────────────────────────────────────────
# TAB 5: HORÓMETROS EXCEL
# ─────────────────────────────────────────────
with tabs[4]:
    st.subheader("📂 Carga de horómetros desde Excel")
    st.caption("Sube el archivo Control_Horómetro_TMF_Rev10.xlsm. Los valores negativos o absurdos se corrigen automáticamente.")

    archivo = st.file_uploader("Seleccionar archivo Excel (.xlsm / .xlsx)", type=["xlsm","xlsx"], key="upload_horo")

    if archivo is not None:
        try:
            df_preview = leer_excel_horometros(archivo)
            if df_preview.empty:
                st.warning("No se encontraron equipos GET en el archivo.")
            else:
                if "tiene_error" in df_preview.columns:
                    errores_excel = df_preview[df_preview["tiene_error"] == True]
                    if not errores_excel.empty:
                        st.warning(f"⚠️ Se detectaron y corrigieron valores inválidos en {len(errores_excel)} equipos: **{', '.join(errores_excel['equipo'].tolist())}**.")

                st.success(f"✅ Se encontraron **{len(df_preview)} equipos** en el Excel.")
                cols_show = ["equipo","fecha","horometro_actual","promedio_7d","promedio_30d","promedio_historico"]
                cols_show = [c for c in cols_show if c in df_preview.columns]
                st.dataframe(df_preview[cols_show], use_container_width=True)

                if st.button("💾 Guardar horómetros en base de datos", type="primary", key="btn_guardar_horo"):
                    n = guardar_horometros(df_preview)
                    st.success(f"✅ {n} equipos actualizados en Supabase.")
                    st.rerun()
        except Exception as e:
            st.error(f"Error leyendo el archivo: {e}")

    st.divider()
    st.subheader("Último horómetro cargado por equipo")
    df_horo_db = cargar_horometros_db()
    if not df_horo_db.empty:
        cols_h = [c for c in ["equipo","fecha","horometro_actual","promedio_7d","promedio_30d","promedio_historico"] if c in df_horo_db.columns]
        st.dataframe(df_horo_db[cols_h].sort_values("equipo"), use_container_width=True)
    else:
        st.info("Sin horómetros cargados aún.")


# ─────────────────────────────────────────────
# TAB 6: REPORTE SEMANAL
# ─────────────────────────────────────────────
with tabs[5]:
    st.subheader("📊 Reporte Semanal")
    df_rep    = cargar_historial(5000)
    ultimos_r = ultimos_por_equipo(df_rep)
    df_dias_r = dias_sin_medicion(df_rep)
    hoy_rep = date.today()
    dias_desde_jueves = (hoy_rep.weekday() - 3) % 7
    jueves_semana = hoy_rep - timedelta(days=dias_desde_jueves)
    semana = jueves_semana.isocalendar()[1]

    st.markdown(f"**Semana {semana} · Período: {jueves_semana.strftime('%d/%m/%Y')} → {(jueves_semana + timedelta(days=6)).strftime('%d/%m/%Y')}**")

    if not ultimos_r.empty:
        miercoles_semana = jueves_semana + timedelta(days=6)
        df_semana = df_rep[
            (pd.to_datetime(df_rep["fecha"], errors="coerce").dt.date >= jueves_semana) &
            (pd.to_datetime(df_rep["fecha"], errors="coerce").dt.date <= miercoles_semana)
        ] if not df_rep.empty else pd.DataFrame()
        equipos_semana = sorted(df_semana["equipo"].unique().tolist()) if not df_semana.empty else []

        st.markdown(f"#### 📅 Semana {semana} — mediciones del {jueves_semana.strftime('%d/%m')} al {miercoles_semana.strftime('%d/%m/%Y')}")
        if equipos_semana:
            st.success(f"✅ **{len(equipos_semana)} equipos medidos esta semana:** {', '.join(equipos_semana)}")
            ultimos_semana = ultimos_por_equipo(df_semana)
            for _, row in ultimos_semana.iterrows():
                estado = str(row.get("estado",""))
                icon   = COLOR_ESTADO.get(estado,"⚪")
                mm     = f"{row['mm_usada']:.1f} mm" if pd.notna(row.get("mm_usada")) else "—"
                pct_val = row.get("condicion_pct")
                pct_html2 = badge_pct_html(pct_val) if pd.notna(pct_val) else "—"
                st_style = card_style(estado)
                bg, border, txt = _ESTADO_CARD.get(estado, (BG_NODAT, "#D3D1C7", "#444441"))
                st.markdown(
                    f'<div style="{st_style}">'
                    f'{dot_estado(estado)}<b style="color:{txt};">Equipo {row["equipo"]}</b>'
                    f' &nbsp;|&nbsp; <span style="color:{txt};">{icon} {estado}</span>'
                    f' &nbsp;|&nbsp; <span style="color:{txt};">{mm}</span>'
                    f' &nbsp;|&nbsp; Desgaste: {pct_html2}'
                    f'</div>', unsafe_allow_html=True)
        else:
            st.warning("⚠️ Sin mediciones registradas en la semana en curso.")

        st.divider()
        st.markdown("#### 📊 Estado actual de flota (último registro por equipo)")
        for _, row in ultimos_r.iterrows():
            estado = str(row.get("estado",""))
            icon   = COLOR_ESTADO.get(estado,"⚪")
            mm     = f"{row['mm_usada']:.1f} mm" if pd.notna(row.get("mm_usada")) else "—"
            bg, border, txt = _ESTADO_CARD.get(estado, (BG_NODAT, "#D3D1C7", "#444441"))

            fecha_med = row.get("fecha")
            dias_sin_html = "—"
            if fecha_med:
                dias_sin = (date.today() - pd.to_datetime(fecha_med).date()).days
                dias_sin_html = badge_dias_html(dias_sin)

            pct_val = row.get("condicion_pct")
            pct_html = badge_pct_html(pct_val) if pd.notna(pct_val) else "—"

            st.markdown(
                f'<div style="{card_style(estado)}">'
                f'{dot_estado(estado)}<b style="color:{txt};">Equipo {row["equipo"]}</b>'
                f' &nbsp;|&nbsp; <span style="color:{txt};">{icon} {estado}</span>'
                f' &nbsp;|&nbsp; <span style="color:{txt};">{mm}</span>'
                f' &nbsp;|&nbsp; Desgaste: {pct_html}'
                f' &nbsp;|&nbsp; {dias_sin_html}'
                f'</div>', unsafe_allow_html=True)

        st.divider()
        st.markdown("#### Equipos sin medir (> 7 días)")
        sin_medir_r = df_dias_r[df_dias_r["dias_sin_medir"].notna() & (df_dias_r["dias_sin_medir"] > 7)]
        if not sin_medir_r.empty:
            for _, r in sin_medir_r.iterrows():
                d = int(r["dias_sin_medir"])
                st.markdown(
                    f'<div style="{card_style("MEDIO") if d <= 14 else card_style("CRÍTICO")}">'
                    f'{dot_estado("MEDIO" if d <= 14 else "CRÍTICO")}'
                    f'<b>Equipo {r["equipo"]}</b>'
                    f' &nbsp;|&nbsp; Última medición: {r["ultima_medicion"]}'
                    f' &nbsp;|&nbsp; {badge_dias_html(d)}'
                    f'</div>', unsafe_allow_html=True)
        else:
            st.success("✅ Todos los equipos medidos en los últimos 7 días.")

        st.divider()
        st.markdown("#### Enviar reporte por correo")
        dest_input = st.text_input("Destinatarios (separados por coma)", value="Pablo.cortes2@teck.com", key="dest_mail")

        col_prev, col_env = st.columns(2)
        with col_prev:
            if st.button("👁️ Ver preview", key="btn_prev"):
                html_prev = generar_html_reporte(ultimos_r, df_dias_r)
                st.components.v1.html(html_prev, height=600, scrolling=True)
        with col_env:
            if st.button("📧 Enviar reporte", type="primary", key="btn_mail"):
                dests = [d.strip() for d in dest_input.split(",") if d.strip()]
                if not dests:
                    st.error("Ingresa al menos un destinatario.")
                else:
                    cols_xl = [c for c in ["equipo","fecha","mm_usada","condicion_pct","estado","tasa_mm_h","dias_a_critico"] if c in ultimos_r.columns]
                    ok, msg = enviar_correo(dests, generar_html_reporte(ultimos_r, df_dias_r),
                                           excel_bytes(ultimos_r[cols_xl]),
                                           f"GET_Wear_S{semana}_{date.today()}.xlsx")
                    if ok: st.success(f"✅ {msg}")
                    else:  st.error(f"❌ {msg}")

        st.download_button("⬇️ Descargar Excel reporte",
            data=excel_bytes(ultimos_r),
            file_name=f"reporte_semana_{semana}.xlsx")
    else:
        st.info("Sin datos para generar reporte.")


# ─────────────────────────────────────────────
# TAB 7: REPORTE EJECUTIVO
# ─────────────────────────────────────────────
with tabs[6]:
    st.subheader("📋 Resumen de Equipos")
    st.caption("Estado actual de todos los equipos con última medición, días sin medir y condición.")

    df_res = cargar_historial(5000)
    df_horo_res = cargar_horometros_db()

    if not df_res.empty:
        hoy = date.today()
        filas = []
        for eq in EQUIPOS:
            df_eq = df_res[df_res["equipo"] == eq]
            if df_eq.empty:
                filas.append({
                    "Equipo": eq, "Última medición": "—", "Días sin medir": "—",
                    "mm usada": "—", "Desgaste %": "—", "Estado": "SIN DATOS",
                    "Horómetro medición": "—", "Horómetro Excel": "—",
                    "Técnico": "—",
                })
                continue

            last = df_eq.sort_values("fecha", ascending=False).iloc[0]
            ultima_fecha = str(last.get("fecha",""))[:10]
            try:
                dias_sin = (hoy - pd.to_datetime(ultima_fecha).date()).days
                dias_txt = str(dias_sin)
            except:
                dias_txt = "—"

            h_excel = "—"
            if not df_horo_res.empty:
                row_h = df_horo_res[df_horo_res["equipo"] == eq]
                if not row_h.empty:
                    h_excel = f"{row_h.iloc[0]['horometro_actual']:,.0f}"

            estado = str(last.get("estado","—"))
            mm = f"{last['mm_usada']:.1f}" if pd.notna(last.get("mm_usada")) else "—"
            pct = f"{last['condicion_pct']:.1f}%" if pd.notna(last.get("condicion_pct")) else "—"

            filas.append({
                "Equipo": eq,
                "Última medición": ultima_fecha,
                "Días sin medir": dias_txt,
                "mm usada": mm,
                "Desgaste %": pct,
                "Estado": estado,
                "Horómetro medición": f"{last['horometro']:,.0f}" if pd.notna(last.get("horometro")) else "—",
                "Horómetro Excel": h_excel,
                "Técnico": str(last.get("usuario","—")),
            })

        df_tabla = pd.DataFrame(filas)

        def color_estado_row(val):
            m = {
                "OK":       f"background-color:{BG_OK};color:{TEXT_OK}",
                "MEDIO":    f"background-color:{BG_MEDIO};color:{TEXT_MEDIO}",
                "ALTO":     f"background-color:{BG_ALTO};color:{TEXT_ALTO}",
                "CRÍTICO":  f"background-color:{BG_CRIT};color:{TEXT_CRIT}",
                "SIN DATOS":f"color:#888780",
            }
            return m.get(val,"")

        def color_dias_row(val):
            try:
                d = int(val)
                if d <= 10:  return f"color:{DIAS_OK_FG};font-weight:bold"
                if d <= 14:  return f"color:{DIAS_MD_FG};font-weight:bold"
                return f"color:{DIAS_CR_BG};font-weight:bold"
            except:
                return "color:#888780"

        st.dataframe(
            df_tabla.style
                .map(color_estado_row, subset=["Estado"])
                .map(color_dias_row, subset=["Días sin medir"]),
            use_container_width=True,
            height=600
        )

        st.divider()
        k1, k2, k3, k4, k5 = st.columns(5)
        sin_datos = len([f for f in filas if f["Estado"]=="SIN DATOS"])
        ok_count  = len([f for f in filas if f["Estado"]=="OK"])
        med_count = len([f for f in filas if f["Estado"]=="MEDIO"])
        alt_count = len([f for f in filas if f["Estado"]=="ALTO"])
        crit_count= len([f for f in filas if f["Estado"]=="CRÍTICO"])
        k1.metric("🟢 OK", ok_count)
        k2.metric("🟡 Monitoreo", med_count)
        k3.metric("🟠 Programar", alt_count)
        k4.metric("🔴 Crítico", crit_count)
        k5.metric("⚫ Sin datos", sin_datos)

        criticos = [f["Equipo"] for f in filas if f["Estado"] in ("CRÍTICO","ALTO")]
        atrasados = [f for f in filas if f["Días sin medir"] not in ("—",) and int(f["Días sin medir"]) > 7]

        if criticos:
            st.error(f"🔴 **Requieren acción inmediata:** {', '.join(criticos)}")
        if atrasados:
            txt = ", ".join([f"Eq {f['Equipo']} ({f['Días sin medir']} días)" for f in atrasados])
            st.warning(f"⚠️ **Sin medir hace más de 7 días:** {txt}")
        if not criticos and not atrasados:
            st.success("✅ Todos los equipos al día y en condición normal.")

        st.divider()
        st.markdown("#### Generar Reporte Ejecutivo Word (.docx)")
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            periodo_input = st.text_input("Período del reporte", value=f"{jueves_semana.strftime('%d/%m/%Y') if 'jueves_semana' in dir() else date.today().strftime('%d/%m/%Y')} — {date.today().strftime('%d/%m/%Y')}", key="periodo_ej")
        with col_d2:
            if st.button("📄 Generar DOCX Ejecutivo", type="primary", key="btn_docx_ej"):
                flota_list = []
                for f in filas:
                    if f["Estado"] not in ("SIN DATOS",):
                        try:
                            mm_v = float(str(f["mm usada"]).replace(",","."))
                            pct_v = float(str(f["Desgaste %"]).replace("%","").replace(",","."))
                            dias_c_v = None
                            row_u = df_res[df_res["equipo"] == f["Equipo"]].sort_values("fecha", ascending=False)
                            if not row_u.empty:
                                dias_c_v = row_u.iloc[0].get("dias_a_critico")
                                if pd.notna(dias_c_v): dias_c_v = float(dias_c_v)
                                else: dias_c_v = None
                            tipo = "Motoniveladora" if f["Equipo"] in EQUIPOS_MOTONIVELADORA else "Dozer/Wheeldozer"
                            flota_list.append({
                                "equipo": f["Equipo"], "tipo": tipo,
                                "mm": mm_v, "pct": pct_v,
                                "estado": f["Estado"], "d_crit": dias_c_v,
                            })
                        except: pass
                if flota_list:
                    hoy_ej = date.today()
                    dias_j = (hoy_ej.weekday() - 3) % 7
                    jue_ej = hoy_ej - timedelta(days=dias_j)
                    sem_ej = jue_ej.isocalendar()[1]
                    anio_ej = hoy_ej.year
                    docx_bytes = generar_reporte_ejecutivo_docx(
                        flota_list, periodo_input,
                        hoy_ej.strftime("%d/%m/%Y"), sem_ej, anio_ej
                    )
                    st.download_button(
                        "⬇️ Descargar Reporte Ejecutivo",
                        data=docx_bytes,
                        file_name=f"Reporte_Ejecutivo_S{sem_ej}_{hoy_ej}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="dl_docx_ej"
                    )
                    st.success("✅ Reporte generado. Haz clic en el botón de descarga.")
                else:
                    st.warning("No hay datos suficientes para generar el reporte.")
    else:
        st.info("Sin mediciones registradas.")


# ─────────────────────────────────────────────
# TAB PROYECCIÓN DE CAMBIOS — solo administrador
# ─────────────────────────────────────────────
if admin_ok and TAB_PROYECCION is not None:
    with tabs[TAB_PROYECCION]:
        st.subheader("📆 Proyección de Fechas de Cambio")
        st.caption(f"Calcula la fecha estimada de cambio de cuchilla para cada equipo. Requiere mínimo **{MIN_MEDS_PROYECCION} mediciones** en el ciclo actual.")
        st.warning("⚠️ **Modo prueba** — Las proyecciones son estimativas basadas en tasa de desgaste histórica. Validar con medición en terreno.")

        st.divider()

        if st.button("🔄 Calcular proyecciones de toda la flota", type="primary", key="btn_proy_all"):
            st.session_state["proy_calculada"] = True
            resultados = {}
            prog = st.progress(0)
            for i, eq in enumerate(EQUIPOS):
                resultados[eq] = proyectar_fecha_cambio(eq)
                prog.progress((i+1)/len(EQUIPOS))
            st.session_state["proy_resultados"] = resultados
            prog.empty()

        if st.session_state.get("proy_calculada") and "proy_resultados" in st.session_state:
            resultados = st.session_state["proy_resultados"]

            con_proyeccion = {eq:r for eq,r in resultados.items() if r["ok"]}
            sin_proyeccion = {eq:r for eq,r in resultados.items() if not r["ok"]}

            urgentes  = [eq for eq,r in con_proyeccion.items() if r.get("dias_restantes",999) <= 7]
            proximos  = [eq for eq,r in con_proyeccion.items() if 7 < r.get("dias_restantes",999) <= 30]
            normales  = [eq for eq,r in con_proyeccion.items() if r.get("dias_restantes",999) > 30]

            k1,k2,k3,k4 = st.columns(4)
            k1.metric("🔴 Cambio urgente (≤7 días)",    len(urgentes))
            k2.metric("🟠 Cambio próximo (8-30 días)",   len(proximos))
            k3.metric("🟢 Con margen (>30 días)",        len(normales))
            k4.metric("⚫ Sin datos suficientes",         len(sin_proyeccion))

            st.divider()
            st.markdown("### Tabla de proyección por equipo")

            filas_proy = []
            for eq in EQUIPOS:
                r = resultados.get(eq, {})
                if r.get("ok"):
                    dias = r.get("dias_restantes", 0)
                    if dias <= 7:     urg = "🔴 URGENTE"
                    elif dias <= 14:  urg = "🟠 ESTA SEMANA"
                    elif dias <= 30:  urg = "🟡 ESTE MES"
                    else:             urg = "🟢 CON MARGEN"
                    filas_proy.append({
                        "Equipo": eq,
                        "mm actual": f"{r['mm_actual']:.1f}",
                        "mm crítico": f"{r['mm_critico']:.0f}",
                        "Tasa mm/h": f"{r['tasa']:.4f}",
                        "h/día (Excel)": f"{r['h_dia']:.1f}" if r.get("h_dia") else "—",
                        "Horas restantes": f"{r['horas_restantes']:.0f} h",
                        "Días restantes": f"{r['dias_restantes']:.1f}",
                        "📅 Fecha estimada cambio": r["fecha_cambio"].strftime("%d/%m/%Y"),
                        "Mediciones ciclo": r["meds"],
                        "Confianza": r["confianza"],
                        "Prioridad": urg,
                    })
                else:
                    filas_proy.append({
                        "Equipo": eq,
                        "mm actual": "—", "mm crítico": "—", "Tasa mm/h": "—",
                        "h/día (Excel)": "—", "Horas restantes": "—", "Días restantes": "—",
                        "📅 Fecha estimada cambio": "Sin datos",
                        "Mediciones ciclo": r.get("meds", 0),
                        "Confianza": "—",
                        "Prioridad": f"⚫ {r.get('error','Sin datos')[:40]}",
                    })

            df_proy = pd.DataFrame(filas_proy)

            def color_prioridad(val):
                if "URGENTE"    in str(val): return f"background-color:{BG_CRIT};color:{TEXT_CRIT}"
                if "ESTA SEMANA"in str(val): return f"background-color:{BG_ALTO};color:{TEXT_ALTO}"
                if "ESTE MES"   in str(val): return f"background-color:{BG_MEDIO};color:{TEXT_MEDIO}"
                if "CON MARGEN" in str(val): return f"background-color:{BG_OK};color:{TEXT_OK}"
                return "color:#888780"

            st.dataframe(
                df_proy.style.map(color_prioridad, subset=["Prioridad"]),
                use_container_width=True
            )

            st.download_button(
                "⬇️ Descargar proyección Excel",
                data=excel_bytes(df_proy),
                file_name=f"proyeccion_cambios_{date.today()}.xlsx",
                key="dl_proy"
            )

            st.divider()
            st.markdown("### Detalle por equipo")
            eq_sel_proy = st.selectbox("Seleccionar equipo para ver detalle", EQUIPOS, key="eq_proy_det")
            r_det = resultados.get(eq_sel_proy, {})

            if r_det.get("ok"):
                ca, cb, cc, cd = st.columns(4)
                ca.metric("mm actual",     f"{r_det['mm_actual']:.1f} mm")
                cb.metric("mm crítico",    f"{r_det['mm_critico']:.0f} mm")
                cc.metric("Tasa desgaste", f"{r_det['tasa']:.4f} mm/h")
                cd.metric("h/día equipo",  f"{r_det['h_dia']:.1f}" if r_det.get('h_dia') else "—")

                ce, cf, cg, ch = st.columns(4)
                ce.metric("Horas restantes",   f"{r_det['horas_restantes']:.0f} h")
                cf.metric("Días restantes",    f"{r_det['dias_restantes']:.1f} días")
                cg.metric("📅 Fecha estimada", r_det["fecha_cambio"].strftime("%d/%m/%Y"))
                ch.metric("Confianza",         r_det["confianza"])

                st.info(f"📊 Basado en **{r_det['meds']} mediciones** del ciclo actual · Última medición: {r_det.get('fecha_ultima_med','—')}")

                if r_det["confianza"] == "Media":
                    st.warning(f"⚠️ Confianza Media — se tienen {r_det['meds']} mediciones. Con 5 o más la proyección es más precisa.")
                elif r_det["confianza"] == "Alta":
                    st.success(f"✅ Confianza Alta — {r_det['meds']} mediciones en el ciclo actual.")
            else:
                st.error(f"No se puede proyectar para equipo {eq_sel_proy}: {r_det.get('error','Sin datos')}")
                st.info(f"Mediciones actuales: {r_det.get('meds',0)} de {MIN_MEDS_PROYECCION} requeridas.")
        else:
            st.info("Haz clic en **Calcular proyecciones** para ver las fechas estimadas de cambio.")


# ─────────────────────────────────────────────
# TAB ACCESOS — solo administrador
# ─────────────────────────────────────────────
if admin_ok and TAB_ACCESOS is not None:
    with tabs[TAB_ACCESOS]:
        st.subheader("👥 Registro de Accesos y Actividad")
        st.caption("Muestra quién ha ingresado mediciones, cambios y desde qué fecha. Solo visible para administradores.")

        col_a1, col_a2 = st.columns(2)

        with col_a1:
            st.markdown("#### 📏 Técnicos que han ingresado mediciones")
            df_acc = cargar_historial(5000)
            if not df_acc.empty and "usuario" in df_acc.columns:
                df_acc["fecha_dt"] = pd.to_datetime(df_acc["fecha"], errors="coerce")
                resumen_usuarios = df_acc.groupby("usuario").agg(
                    mediciones=("id","count"),
                    primera=("fecha","min"),
                    ultima=("fecha","max"),
                    equipos=("equipo", lambda x: ", ".join(sorted(x.unique())))
                ).reset_index().sort_values("ultima", ascending=False)
                resumen_usuarios.columns = ["Técnico","Mediciones","Primera","Última","Equipos"]
                st.dataframe(resumen_usuarios, use_container_width=True)
            else:
                st.info("Sin registros.")

        with col_a2:
            st.markdown("#### 🔄 Técnicos que han registrado cambios")
            df_cam_acc = cargar_cambios()
            if not df_cam_acc.empty:
                resumen_cambios = df_cam_acc.groupby("usuario").agg(
                    cambios=("id","count"),
                    primera=("fecha","min"),
                    ultima=("fecha","max"),
                ).reset_index().sort_values("ultima", ascending=False)
                resumen_cambios.columns = ["Registrado por","Cambios","Primera","Última"]
                st.dataframe(resumen_cambios, use_container_width=True)

                st.markdown("#### 👷 Técnicos participantes en cambios")
                tec_list = []
                for _, r in df_cam_acc.iterrows():
                    if r.get("tecnico_1"): tec_list.append({"tecnico": r["tecnico_1"], "equipo": r["equipo"], "fecha": r["fecha"]})
                    if r.get("tecnico_2") and str(r.get("tecnico_2")).strip(): tec_list.append({"tecnico": r["tecnico_2"], "equipo": r["equipo"], "fecha": r["fecha"]})
                if tec_list:
                    df_tec = pd.DataFrame(tec_list)
                    resumen_tec = df_tec.groupby("tecnico").agg(
                        participaciones=("equipo","count"),
                        ultima=("fecha","max"),
                        equipos=("equipo", lambda x: ", ".join(sorted(x.unique())))
                    ).reset_index().sort_values("participaciones", ascending=False)
                    resumen_tec.columns = ["Técnico","Participaciones","Última","Equipos"]
                    st.dataframe(resumen_tec, use_container_width=True)
            else:
                st.info("Sin cambios registrados.")

        st.divider()
        st.markdown("#### 📅 Actividad reciente (últimas 50 acciones)")
        df_rec = cargar_historial(50)
        if not df_rec.empty:
            cols_rec = [c for c in ["fecha","equipo","usuario","estado","mm_usada","horometro"] if c in df_rec.columns]
            df_rec_show = df_rec[cols_rec].copy()
            df_rec_show.insert(0, "tipo", "📏 Medición")
            df_cam_rec = cargar_cambios(20)
            if not df_cam_rec.empty:
                cols_cam = [c for c in ["fecha","equipo","usuario","horometro"] if c in df_cam_rec.columns]
                df_cam_show = df_cam_rec[cols_cam].copy()
                df_cam_show.insert(0, "tipo", "🔄 Cambio GET")
                df_cam_show["estado"] = "CAMBIO"
                df_cam_show["mm_usada"] = "—"
                df_actividad = pd.concat([df_rec_show, df_cam_show], ignore_index=True)
            else:
                df_actividad = df_rec_show
            df_actividad = df_actividad.sort_values("fecha", ascending=False).head(50)
            st.dataframe(df_actividad, use_container_width=True)

        st.divider()
        st.markdown("#### 📊 Estadísticas generales")
        ka, kb, kc, kd = st.columns(4)
        total_meds = len(cargar_historial(9999)) if not cargar_historial().empty else 0
        total_tec  = df_acc["usuario"].nunique() if not df_acc.empty and "usuario" in df_acc.columns else 0
        total_cam  = len(cargar_cambios()) if not cargar_cambios().empty else 0
        primera_fecha = df_acc["fecha"].min() if not df_acc.empty else "—"
        ka.metric("Total mediciones", total_meds)
        kb.metric("Técnicos activos", total_tec)
        kc.metric("Cambios GET registrados", total_cam)
